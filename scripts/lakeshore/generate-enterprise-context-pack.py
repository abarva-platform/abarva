#!/usr/bin/env python3
"""
LAKESHORE_ENTERPRISE_CONTEXT_LOAD_V1 — synthetic enterprise context pack generator.

Generates a best-in-class, fully synthetic enterprise transformation data room for
"Lakeshore Holdings" across 12 context domains. Every artifact is watermarked
SYNTHETIC and carries loader metadata (context_domain, source_owner, source_system,
source_date, sensitivity, synthetic_flag, evidence_usable_flag, loader_route).

Outputs to: docs/build/lakeshore-enterprise-context/

This generator is deterministic (seeded) so runs are reproducible.
"""
from __future__ import annotations
import csv, json, hashlib, io, os, random, textwrap, zipfile, datetime as dt
from pathlib import Path

# ----------------------------------------------------------------------------
# Constants / configuration
# ----------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "build" / "lakeshore-enterprise-context"
SRC = OUT / "source"
PACK_ID = "LAKESHORE_ENTERPRISE_CONTEXT_LOAD_V1"
GEN_DATE = "2026-06-06"
TENANT_KEY = "lakeshore-holdings"
CLIENT_NAME = "Lakeshore Holdings"
WATERMARK = "SYNTHETIC — LAKESHORE PILOT — NOT REAL DATA — ABARVA AI"
WATERMARK_SHORT = "SYNTHETIC / LAKESHORE PILOT / NOT REAL DATA"
random.seed(20260606)

DOMAINS = {
    "00_manifest": "manifest",
    "01_enterprise_profile": "enterprise_profile",
    "02_org_decision_rights": "org_decision_rights",
    "03_strategy_initiatives": "strategy_initiatives",
    "04_finance_performance": "finance_performance",
    "05_treasury_kyriba": "treasury_kyriba",
    "06_it_systems_architecture": "it_systems_architecture",
    "07_data_analytics_reporting": "data_analytics_reporting",
    "08_operations_business_process": "operations_business_process",
    "09_servicenow_support_workload": "servicenow_support_workload",
    "10_vendors_contracts_source": "vendors_contracts_source",
    "11_risk_controls_responsible_ai": "risk_controls_responsible_ai",
    "12_ai_use_cases_moves": "ai_use_cases_moves",
}

# Enterprise canon — reused across every file for internal consistency
BUSINESS_UNITS = [
    ("LIC", "Lakeshore Industrial Components", 3120, "Industrial Manufacturing"),
    ("LCB", "Lakeshore Consumer Brands", 2180, "Consumer Products"),
    ("LLD", "Lakeshore Logistics & Distribution", 1490, "Transportation & Logistics"),
    ("LFS", "Lakeshore Financial Services", 980, "Specialty Finance"),
    ("LHS", "Lakeshore Health Supplies", 640, "Medical Distribution"),
]
GEOS = [
    ("NA-US", "United States", "North America", 4980),
    ("NA-CA", "Canada", "North America", 610),
    ("NA-MX", "Mexico", "North America", 520),
    ("EU-UK", "United Kingdom", "EMEA", 690),
    ("EU-DE", "Germany", "EMEA", 540),
    ("EU-PL", "Poland", "EMEA", 230),
    ("AP-SG", "Singapore", "APAC", 360),
    ("AP-IN", "India", "APAC", 280),
    ("AP-AU", "Australia", "APAC", 210),
    ("LA-BR", "Brazil", "LATAM", 190),
]
SYSTEMS = [
    "SAP ECC 6.0", "SAP S/4HANA", "Workday HCM", "Coupa", "Kyriba TMS",
    "ServiceNow ITSM", "Snowflake", "Power BI", "Azure Active Directory",
    "Oracle Hyperion", "Blackline", "Salesforce", "OneStream", "HighRadius",
]
BANKS = ["JPMorgan", "Citi", "HSBC", "BNP Paribas", "Standard Chartered",
         "Bank of America", "Deutsche Bank", "DBS", "Santander", "Wells Fargo"]
EXECS = [
    ("Maria Donnelly", "Group CEO"), ("Raymond Okafor", "Group CFO"),
    ("Priya Natarajan", "Group CIO"), ("Tomas Halvorsen", "Group CPO"),
    ("Elena Vasquez", "Group Treasurer"), ("David Chen", "Chief Risk Officer"),
    ("Aisha Bello", "Chief Data Officer"), ("Marcus Reilly", "Chief Transformation Officer"),
    ("Sofia Marchetti", "VP FP&A"), ("Liam O'Sullivan", "VP IT Architecture"),
]
SI_PARTNERS = ["Accenture", "Deloitte", "Capgemini", "Infosys", "TCS", "PwC", "EY", "Cognizant"]

# Per-file metadata registry — populated by builders, consumed by manifest builder
REGISTRY: list[dict] = []

# ----------------------------------------------------------------------------
# Watermark helpers
# ----------------------------------------------------------------------------
def _route_for(ext: str) -> str:
    structured = {"csv", "json", "jsonl", "yaml", "yml"}
    if ext in structured:
        return "setup-admin/context-layer/csv-upload (stage_and_process)"
    return "setup-admin/context-layer/bulk-upload (stage_and_enqueue -> service-bus -> worker)"

def register(path: Path, *, domain_folder: str, owner: str, system: str,
             date: str, sensitivity: str, evidence_usable: bool,
             title: str, description: str):
    rel = path.relative_to(OUT).as_posix()
    ext = path.suffix.lower().lstrip(".")
    data = path.read_bytes()
    REGISTRY.append({
        "file_path": rel,
        "zip_path": path.relative_to(SRC.parent).as_posix() if SRC in path.parents else rel,
        "file_name": path.name,
        "file_type": ext,
        "bytes": len(data),
        "sha256": hashlib.sha256(data).hexdigest(),
        "title": title,
        "description": description,
        "context_domain": DOMAINS.get(domain_folder, domain_folder),
        "domain_folder": domain_folder,
        "source_owner": owner,
        "source_system": system,
        "source_date": date,
        "sensitivity": sensitivity,
        "synthetic_flag": True,
        "synthetic_label": WATERMARK_SHORT,
        "evidence_usable_flag": evidence_usable,
        "loader_route": _route_for(ext),
    })

# ----------------------------------------------------------------------------
# File-type writers (each auto-applies watermark)
# ----------------------------------------------------------------------------
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

def write_xlsx(path: Path, sheets: dict[str, tuple[list[str], list[list]]], *, title: str, meta: dict):
    wb = Workbook()
    wb.remove(wb.active)
    banner_fill = PatternFill("solid", fgColor="1F2A44")
    hdr_fill = PatternFill("solid", fgColor="D8DEE9")
    for sheet_name, (headers, rows) in sheets.items():
        ws = wb.create_sheet(sheet_name[:31])
        ws.append([f"{WATERMARK}  |  {title}  |  context_domain={meta['domain']}  |  generated={GEN_DATE}"])
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(1, len(headers)))
        c = ws.cell(row=1, column=1); c.font = Font(bold=True, color="FFFFFF", size=9); c.fill = banner_fill
        ws.append(headers)
        for col in range(1, len(headers) + 1):
            hc = ws.cell(row=2, column=col); hc.font = Font(bold=True); hc.fill = hdr_fill
        for r in rows:
            ws.append(r)
        for col in range(1, len(headers) + 1):
            letter = get_column_letter(col)
            maxlen = max([len(str(headers[col-1]))] + [len(str(r[col-1])) for r in rows[:200] if col-1 < len(r)] + [12])
            ws.column_dimensions[letter].width = min(48, maxlen + 2)
        ws.freeze_panes = "A3"
    props = wb.properties
    props.creator = "AbarVa Synthetic Generator"
    props.title = f"{title} [{WATERMARK_SHORT}]"
    props.subject = WATERMARK_SHORT
    props.keywords = f"synthetic;lakeshore-pilot;{meta['domain']}"
    props.description = f"{WATERMARK}. Owner={meta['owner']}; System={meta['system']}."
    wb.save(path)

from reportlab.lib.pagesizes import letter as PDF_LETTER
from reportlab.lib.units import inch
from reportlab.lib import colors as rl_colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                                PageBreak, ListFlowable, ListItem)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT

def _pdf_watermark(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica-Bold", 7)
    canvas.setFillColor(rl_colors.HexColor("#8a93a6"))
    canvas.drawString(0.6*inch, 0.4*inch, WATERMARK)
    canvas.drawRightString(7.9*inch, 0.4*inch, f"{PACK_ID} | page {doc.page}")
    canvas.setFont("Helvetica-Bold", 44)
    canvas.setFillColor(rl_colors.Color(0.86, 0.88, 0.92, alpha=0.18))
    canvas.translate(4.25*inch, 5.5*inch); canvas.rotate(40)
    canvas.drawCentredString(0, 0, "SYNTHETIC")
    canvas.restoreState()

def write_pdf(path: Path, title: str, subtitle: str, sections: list[tuple[str, list]], *, meta: dict):
    """sections = list of (heading, body) where body is list of str (paragraph) or
    ('table', headers, rows) or ('bullets', [..])."""
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], textColor=rl_colors.HexColor("#0c1a3a"), fontSize=18)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=rl_colors.HexColor("#1F2A44"), fontSize=12)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=9.5, leading=13, alignment=TA_LEFT)
    eyebrow = ParagraphStyle("eyebrow", parent=styles["BodyText"], fontSize=7.5,
                             textColor=rl_colors.HexColor("#6b7280"), spaceAfter=2)
    doc = SimpleDocTemplate(str(path), pagesize=PDF_LETTER, topMargin=0.7*inch,
                            bottomMargin=0.7*inch, leftMargin=0.7*inch, rightMargin=0.7*inch,
                            title=f"{title} [{WATERMARK_SHORT}]", author="AbarVa Synthetic Generator",
                            subject=WATERMARK_SHORT)
    flow = []
    flow.append(Paragraph(f"LAKESHORE HOLDINGS · {meta['domain'].upper().replace('_',' ')} · {WATERMARK_SHORT}", eyebrow))
    flow.append(Paragraph(title, h1))
    if subtitle:
        flow.append(Paragraph(subtitle, body))
    flow.append(Spacer(1, 0.12*inch))
    flow.append(Paragraph(f"<b>Owner:</b> {meta['owner']} &nbsp;·&nbsp; <b>Source system:</b> {meta['system']} "
                          f"&nbsp;·&nbsp; <b>As of:</b> {meta['date']} &nbsp;·&nbsp; <b>Sensitivity:</b> {meta['sensitivity']}", eyebrow))
    flow.append(Spacer(1, 0.12*inch))
    for heading, blocks in sections:
        if heading:
            flow.append(Paragraph(heading, h2))
        for b in blocks:
            if isinstance(b, tuple) and b[0] == "table":
                _, headers, rows = b
                tbl_data = [headers] + rows
                t = Table(tbl_data, repeatRows=1, hAlign="LEFT")
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0,0), (-1,0), rl_colors.HexColor("#1F2A44")),
                    ("TEXTCOLOR", (0,0), (-1,0), rl_colors.white),
                    ("FONTSIZE", (0,0), (-1,-1), 7.5),
                    ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                    ("GRID", (0,0), (-1,-1), 0.4, rl_colors.HexColor("#c9ccd6")),
                    ("ROWBACKGROUNDS", (0,1), (-1,-1), [rl_colors.white, rl_colors.HexColor("#f3f4f7")]),
                    ("VALIGN", (0,0), (-1,-1), "TOP"),
                ]))
                flow.append(t); flow.append(Spacer(1, 0.1*inch))
            elif isinstance(b, tuple) and b[0] == "bullets":
                items = [ListItem(Paragraph(x, body), leftIndent=10) for x in b[1]]
                flow.append(ListFlowable(items, bulletType="bullet", start="•")); flow.append(Spacer(1, 0.06*inch))
            else:
                flow.append(Paragraph(b, body)); flow.append(Spacer(1, 0.05*inch))
    doc.build(flow, onFirstPage=_pdf_watermark, onLaterPages=_pdf_watermark)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def write_docx(path: Path, title: str, sections: list[tuple[str, list]], *, meta: dict):
    d = Document()
    cp = d.core_properties
    cp.title = f"{title} [{WATERMARK_SHORT}]"; cp.author = "AbarVa Synthetic Generator"
    cp.comments = WATERMARK; cp.category = meta["domain"]; cp.keywords = "synthetic;lakeshore-pilot"
    sec = d.sections[0]
    header = sec.header.paragraphs[0]
    header.text = WATERMARK
    header.runs[0].font.size = Pt(7); header.runs[0].font.color.rgb = RGBColor(0x8a, 0x93, 0xa6)
    footer = sec.footer.paragraphs[0]
    footer.text = f"{PACK_ID} · {meta['domain']} · {WATERMARK_SHORT}"
    footer.runs[0].font.size = Pt(7)
    t = d.add_heading(title, level=0)
    sub = d.add_paragraph(f"Lakeshore Holdings · Owner: {meta['owner']} · Source: {meta['system']} · "
                          f"As of: {meta['date']} · Sensitivity: {meta['sensitivity']}")
    sub.runs[0].font.size = Pt(8); sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    for heading, blocks in sections:
        if heading:
            d.add_heading(heading, level=1)
        for b in blocks:
            if isinstance(b, tuple) and b[0] == "bullets":
                for x in b[1]:
                    d.add_paragraph(x, style="List Bullet")
            elif isinstance(b, tuple) and b[0] == "table":
                _, headers, rows = b
                tbl = d.add_table(rows=1, cols=len(headers)); tbl.style = "Light Grid Accent 1"
                for i, h in enumerate(headers):
                    tbl.rows[0].cells[i].text = str(h)
                for r in rows:
                    cells = tbl.add_row().cells
                    for i, v in enumerate(r):
                        cells[i].text = str(v)
            else:
                d.add_paragraph(b)
    d.save(path)

def write_md(path: Path, title: str, body: str, *, meta: dict):
    front = (f"---\n"
             f"title: \"{title}\"\n"
             f"synthetic: true\n"
             f"synthetic_label: \"{WATERMARK_SHORT}\"\n"
             f"pack: {PACK_ID}\n"
             f"context_domain: {meta['domain']}\n"
             f"source_owner: \"{meta['owner']}\"\n"
             f"source_system: \"{meta['system']}\"\n"
             f"source_date: {meta['date']}\n"
             f"sensitivity: {meta['sensitivity']}\n"
             f"evidence_usable: {str(meta.get('evidence_usable', True)).lower()}\n"
             f"---\n\n")
    banner = f"> **{WATERMARK}**\n\n"
    path.write_text(front + banner + body, encoding="utf-8")

def write_csv(path: Path, headers: list[str], rows: list[list], *, meta: dict):
    with path.open("w", newline="", encoding="utf-8") as f:
        f.write(f"# {WATERMARK} | context_domain={meta['domain']} | owner={meta['owner']} | "
                f"system={meta['system']} | as_of={meta['date']} | sensitivity={meta['sensitivity']}\n")
        w = csv.writer(f)
        w.writerow(headers)
        w.writerows(rows)

def write_jsonl(path: Path, records: list[dict], *, meta: dict):
    with path.open("w", encoding="utf-8") as f:
        f.write(json.dumps({"_meta": True, "synthetic": True, "synthetic_label": WATERMARK_SHORT,
                            "context_domain": meta["domain"], "source_owner": meta["owner"],
                            "source_system": meta["system"], "source_date": meta["date"],
                            "sensitivity": meta["sensitivity"], "pack": PACK_ID}) + "\n")
        for r in records:
            r = {**r, "synthetic_flag": True}
            f.write(json.dumps(r) + "\n")

def write_svg(path: Path, title: str, svg_inner: str, *, meta: dict):
    svg = (f'<svg xmlns="http://www.w3.org/2000/svg" width="1100" height="720" viewBox="0 0 1100 720">\n'
           f'<rect width="1100" height="720" fill="#f5f1eb"/>\n'
           f'<text x="24" y="34" font-family="monospace" font-size="13" fill="#6b7280">{WATERMARK}</text>\n'
           f'<text x="24" y="62" font-family="serif" font-size="26" fill="#0c1a3a" font-weight="700">{title}</text>\n'
           f'<text x="550" y="380" font-family="sans-serif" font-size="120" fill="#0c1a3a" '
           f'opacity="0.06" text-anchor="middle" transform="rotate(-20 550 380)">SYNTHETIC</text>\n'
           f'{svg_inner}\n'
           f'<text x="24" y="708" font-family="monospace" font-size="10" fill="#6b7280">'
           f'{PACK_ID} · {meta["domain"]} · owner={meta["owner"]} · {meta["date"]}</text>\n'
           f'</svg>\n')
    path.write_text(svg, encoding="utf-8")

# ----------------------------------------------------------------------------
# Synthetic data helpers
# ----------------------------------------------------------------------------
def money(x): return f"${x:,.1f}M"
def pct(x): return f"{x:.1f}%"
def date_in(year_start="2024-01-01", year_end="2026-06-01"):
    s = dt.date.fromisoformat(year_start); e = dt.date.fromisoformat(year_end)
    return (s + dt.timedelta(days=random.randint(0, (e - s).days))).isoformat()

print("Generator module loaded.")

# ============================================================================
# DOMAIN BUILDERS
# ============================================================================
def M(folder, owner, system, date=GEN_DATE, sensitivity="Internal", evidence_usable=True):
    return {"domain": DOMAINS[folder], "owner": owner, "system": system, "date": date,
            "sensitivity": sensitivity, "evidence_usable": evidence_usable}

def reg(path, folder, owner, system, title, desc, date=GEN_DATE, sensitivity="Internal", evidence=True):
    register(path, domain_folder=folder, owner=owner, system=system, date=date,
             sensitivity=sensitivity, evidence_usable=evidence, title=title, description=desc)

# ---- 00 manifest ----------------------------------------------------------
def build_manifest_domain():
    f = "00_manifest"
    p = SRC / f / "pack_overview.md"
    write_md(p, "Lakeshore Enterprise Context Pack — Overview",
        textwrap.dedent(f"""
        ## Purpose
        This data room is a **fully synthetic** enterprise transformation context pack for
        **{CLIENT_NAME}** ({TENANT_KEY}). It is designed to exercise the AbarVa Setup Admin
        bulk/ZIP context loader end-to-end: Azure Blob staging, document parsing, Azure
        Postgres chunk commit, and Azure AI Search / vector refresh — and to support
        board-grade Strategic Move generation.

        ## Scope
        12 context domains, ~116 source documents, including high-volume retrieval corpora
        (1,600 ServiceNow incidents, 320 report/workload rows, 120 vendor/contract/rate-card
        rows, 120 risk/control/audit items, 60 AI use-case records).

        ## Company canon
        - **Group revenue baseline:** ${sum(b[2] for b in BUSINESS_UNITS)/1000:.1f}B across {len(BUSINESS_UNITS)} business units and {len(GEOS)} country operations.
        - **Business units:** {", ".join(b[1] for b in BUSINESS_UNITS)}.
        - **Flagship initiative:** Group Treasury modernization on **Kyriba TMS**, plus
          corporate controls uplift, reporting rationalization, and vendor optimization.
        - **Core systems:** {", ".join(SYSTEMS[:8])}.

        ## Loader routes
        - Structured files (CSV/JSON/JSONL/YAML) → `setup-admin/context-layer/csv-upload` (stage_and_process).
        - Rich documents (PDF/DOCX/XLSX/PPTX/SVG) → `setup-admin/context-layer/bulk-upload`
          (stage_and_enqueue → Service Bus → ingestion worker).

        ## Synthetic guarantee
        Every file carries the watermark **"{WATERMARK_SHORT}"** in content and metadata.
        No real customer, employee, bank, or financial data is present.
        """), meta=M(f, "AbarVa Delivery", "AbarVa Generator"))
    reg(p, f, "AbarVa Delivery", "AbarVa Generator", "Pack Overview", "Top-level description of the synthetic enterprise context pack.")

    p = SRC / f / "load_control_plan.md"
    write_md(p, "Load Control Plan",
        textwrap.dedent(f"""
        ## Controlled load sequence
        1. Generate pack + manifest (local generated state).
        2. Upload ZIP through Setup Admin bulk loader → Azure Blob (`context-uploads`).
        3. Parse each file (Document Intelligence / pdf-parse / exceljs / mammoth / Papa).
        4. Commit chunks/facts/evidence to Azure Postgres `enterprise_context_chunks`.
        5. Refresh Azure AI Search `tenant-context-v1` + vector index.
        6. Verify counts by file type, context domain, source system, segment, DB table, search.
        7. Signed-in QA against Lakeshore tenant auth state.
        8. Publish proof index.

        ## Truth states (kept separate)
        | State | Meaning |
        |---|---|
        | local-generated | Files exist on disk + in ZIP |
        | azure-staged | Original bytes in Azure Blob with retrievable path |
        | db-committed | Rows present in Azure Postgres |
        | indexed-searchable | Retrievable from Azure AI Search / vector |
        | signed-in-qa | Answered through the authenticated product |

        ## Tenant scoping
        - App client key: `lakeshore` · canonical/broker key: `{TENANT_KEY}`.
        - All chunks carry `tenant_key={TENANT_KEY}` and `client_id` = clients.id UUID.
        """), meta=M(f, "AbarVa Delivery", "AbarVa Generator"))
    reg(p, f, "AbarVa Delivery", "AbarVa Generator", "Load Control Plan", "Controlled load sequence and truth-state model.")

    p = SRC / f / "glossary_and_taxonomy.md"
    write_md(p, "Glossary & Context Taxonomy",
        "## Domain taxonomy\n\n" + "\n".join(
            f"- **{v}** — `source/{k}/`" for k, v in DOMAINS.items() if k != "00_manifest"
        ) + textwrap.dedent("""

        ## Key terms
        - **TMS** — Treasury Management System (Kyriba).
        - **AMS** — Application Managed Services.
        - **BAFO** — Best And Final Offer (sourcing).
        - **R2R / O2C / P2P** — Record-to-Report / Order-to-Cash / Procure-to-Pay.
        - **SoD** — Segregation of Duties.
        - **RAID** — Risks, Assumptions, Issues, Dependencies.
        - **Value pool** — addressable financial benefit category.
        """), meta=M(f, "Aisha Bello", "AbarVa Generator"))
    reg(p, f, "Aisha Bello", "AbarVa Generator", "Glossary & Taxonomy", "Domain taxonomy and key terms.")

    p = SRC / f / "context_domain_index.csv"
    rows = [[i, k, v, f"source/{k}/"] for i, (k, v) in enumerate(DOMAINS.items())]
    write_csv(p, ["idx", "domain_folder", "context_domain", "path"], rows, meta=M(f, "AbarVa Delivery", "AbarVa Generator"))
    reg(p, f, "AbarVa Delivery", "AbarVa Generator", "Context Domain Index", "Machine-readable list of context domains.")

# ---- 01 enterprise profile ------------------------------------------------
def build_enterprise_profile():
    f = "01_enterprise_profile"; owner = "Marcus Reilly"; sys = "Corporate Strategy"
    p = SRC / f / "corporate_overview.pdf"
    write_pdf(p, "Lakeshore Holdings — Corporate Overview",
        "Diversified industrial and consumer holding company. Synthetic profile for AbarVa context loading.",
        [("Company snapshot", [
            f"Lakeshore Holdings is a diversified holding company operating {len(BUSINESS_UNITS)} business units "
            f"across {len(GEOS)} countries, with a group revenue baseline of approximately "
            f"${sum(b[2] for b in BUSINESS_UNITS)/1000:.1f}B and ~28,400 employees.",
            ("table", ["Business Unit", "Sector", "Revenue ($M)"],
             [[b[1], b[3], f"{b[2]:,}"] for b in BUSINESS_UNITS])]),
         ("Geographic footprint", [
            ("table", ["Region", "Country", "Revenue ($M)"],
             [[g[2], g[1], f"{g[3]:,}"] for g in GEOS])]),
         ("Leadership priorities (FY26-27)", [("bullets", [
            "Modernize group treasury on Kyriba and reduce idle cash by 18-22%.",
            "Rationalize the management reporting estate (320+ reports → target ~140).",
            "Lift corporate controls posture ahead of SOX re-attestation.",
            "Optimize the SI/AMS vendor portfolio and reduce run-rate by 12-15%.",
            "Stand up a governed AI operating model with Responsible AI controls.",
        ])]),
         ("Operating model", [
            "Lakeshore runs a federated operating model: a lean corporate center (finance, treasury, "
            "IT architecture, procurement, risk) with BU-level P&L ownership. Shared services deliver "
            "transactional finance, IT operations, and HR administration from three hubs (Dallas, Krakow, Bengaluru)."]),
        ], meta=M(f, owner, sys))
    reg(p, f, owner, sys, "Corporate Overview", "Executive corporate profile, BUs, geos, priorities, operating model.", sensitivity="Confidential")

    p = SRC / f / "business_units.xlsx"
    rows = [[b[0], b[1], b[3], b[2], round(b[2]*random.uniform(0.08,0.18),1),
             random.randint(2200, 9800), random.choice(["Growth","Optimize","Harvest","Turnaround"])]
            for b in BUSINESS_UNITS]
    write_xlsx(p, {"BusinessUnits": (["bu_code","business_unit","sector","revenue_m","ebit_m","headcount","strategic_posture"], rows)},
               title="Business Units", meta=M(f, owner, sys))
    reg(p, f, owner, sys, "Business Units", "Per-BU revenue, EBIT, headcount, posture.")

    p = SRC / f / "geographic_footprint.xlsx"
    rows = [[g[0], g[1], g[2], g[3], random.randint(400, 6200), random.choice(BANKS), random.choice(["SAP ECC 6.0","SAP S/4HANA"])]
            for g in GEOS]
    write_xlsx(p, {"Geographies": (["geo_code","country","region","revenue_m","headcount","primary_bank","erp"], rows)},
               title="Geographic Footprint", meta=M(f, owner, sys))
    reg(p, f, owner, sys, "Geographic Footprint", "Country revenue, headcount, banking, ERP.")

    p = SRC / f / "revenue_cost_baseline.xlsx"
    cats = ["Revenue","COGS","Gross Profit","SG&A","R&D","EBIT","D&A","EBITDA","Net Working Capital","Capex"]
    rows = []
    for c in cats:
        base = random.uniform(200, 8400)
        rows.append([c]+[round(base*random.uniform(0.9,1.1),1) for _ in range(4)])
    write_xlsx(p, {"Baseline": (["line_item","FY23","FY24","FY25","FY26_plan"], rows)},
               title="Revenue & Cost Baseline", meta=M(f, "Raymond Okafor", "Oracle Hyperion"))
    reg(p, f, "Raymond Okafor", "Oracle Hyperion", "Revenue & Cost Baseline", "Group P&L baseline FY23-FY26.", sensitivity="Confidential")

    p = SRC / f / "leadership_priorities.pdf"
    write_pdf(p, "Leadership Priorities & North-Star Outcomes", "CEO/CFO/CIO priority cascade.",
        [("North-star outcomes", [("bullets", [
            "Cash & liquidity: release $140-180M trapped cash via Kyriba cash visibility.",
            "Control posture: zero material weaknesses at FY26 SOX attestation.",
            "Reporting: 55% reduction in manual close-and-report effort.",
            "Cost: $48-62M run-rate reduction across IT/AMS/treasury operations.",
            "AI: 12 governed AI use cases in production with measured value by FY27."])])],
        meta=M(f, "Maria Donnelly", "Corporate Strategy"))
    reg(p, f, "Maria Donnelly", "Corporate Strategy", "Leadership Priorities", "North-star outcomes and priority cascade.", sensitivity="Confidential")

    p = SRC / f / "operating_model.docx"
    write_docx(p, "Group Operating Model",
        [("Design principles", [("bullets", [
            "Lean corporate center owns policy, architecture, and capital allocation.",
            "BUs own P&L and customer outcomes within group guardrails.",
            "Shared services deliver standardized transactional processes.",
            "A single data & AI platform serves all BUs under federated governance."])]),
         ("Decision forums", [
            "The Group Executive Committee meets monthly; the Transformation Steering Committee "
            "meets bi-weekly; the Treasury Council and Architecture Review Board meet monthly."]),
        ], meta=M(f, "Marcus Reilly", "Corporate Strategy"))
    reg(p, f, "Marcus Reilly", "Corporate Strategy", "Operating Model", "Group operating model design principles and forums.")

    p = SRC / f / "transformation_themes.md"
    write_md(p, "Transformation Themes",
        "## FY26-27 themes\n\n" + "\n".join(f"{i+1}. **{t}**" for i,t in enumerate([
            "Treasury & liquidity modernization (Kyriba)",
            "Corporate controls & SOX uplift",
            "Management reporting rationalization",
            "Vendor & sourcing optimization (SI/AMS)",
            "Data platform consolidation (Snowflake)",
            "Governed enterprise AI & Responsible AI",
        ])), meta=M(f, "Marcus Reilly", "Corporate Strategy"))
    reg(p, f, "Marcus Reilly", "Corporate Strategy", "Transformation Themes", "Six transformation themes for FY26-27.")

    p = SRC / f / "enterprise_fact_sheet.csv"
    rows = [["group_revenue_m", sum(b[2] for b in BUSINESS_UNITS)],
            ["business_units", len(BUSINESS_UNITS)], ["countries", len(GEOS)],
            ["employees", 28400], ["primary_erp", "SAP (ECC + S/4HANA)"],
            ["tms", "Kyriba"], ["data_platform", "Snowflake"], ["fy26_capex_m", 410]]
    write_csv(p, ["metric","value"], rows, meta=M(f, owner, sys))
    reg(p, f, owner, sys, "Enterprise Fact Sheet", "Key enterprise facts in machine-readable form.")

    p = SRC / f / "segment_pnl_summary.xlsx"
    rows = [[b[1], b[2], round(b[2]*random.uniform(0.55,0.7),1), round(b[2]*random.uniform(0.1,0.2),1),
             round(b[2]*random.uniform(0.06,0.14),1)] for b in BUSINESS_UNITS]
    write_xlsx(p, {"SegmentPnL": (["segment","revenue_m","cogs_m","sga_m","ebit_m"], rows)},
               title="Segment P&L Summary", meta=M(f, "Sofia Marchetti", "Oracle Hyperion"))
    reg(p, f, "Sofia Marchetti", "Oracle Hyperion", "Segment P&L Summary", "Per-segment P&L.", sensitivity="Confidential")

print("built: enterprise profile")

# ---- 02 org & decision rights ---------------------------------------------
def build_org():
    f = "02_org_decision_rights"
    def org_rows(area, leaders, n):
        rows=[]
        for i in range(n):
            rows.append([f"{area[:3].upper()}-{i+1:03d}",
                         random.choice(["VP","Sr Director","Director","Sr Manager","Manager"]),
                         random.choice(leaders), area, random.choice([g[1] for g in GEOS]),
                         random.randint(3, 80)])
        return rows
    specs = [
        ("executive_org_chart.xlsx", "Executive", [e[0] for e in EXECS], 14, "Maria Donnelly", "Workday HCM"),
        ("finance_org.xlsx", "Finance", ["Raymond Okafor","Sofia Marchetti"], 22, "Raymond Okafor", "Workday HCM"),
        ("it_org.xlsx", "IT", ["Priya Natarajan","Liam O'Sullivan"], 26, "Priya Natarajan", "Workday HCM"),
        ("procurement_vendor_mgmt_org.xlsx", "Procurement", ["Tomas Halvorsen"], 16, "Tomas Halvorsen", "Coupa"),
        ("shared_services_org.xlsx", "Shared Services", ["Sofia Marchetti"], 20, "Sofia Marchetti", "Workday HCM"),
    ]
    for fn, area, leaders, n, owner, sys in specs:
        p = SRC / f / fn
        write_xlsx(p, {area[:20]: (["position_id","level","reports_to","function","location","span_of_control"], org_rows(area, leaders, n))},
                   title=fn.replace("_"," ").replace(".xlsx","").title(), meta=M(f, owner, sys))
        reg(p, f, owner, sys, fn.replace("_"," ").title(), f"{area} organization structure.")
    p = SRC / f / "steering_committees.pdf"
    write_pdf(p, "Steering Committees & Governance Forums", "Decision forums and cadence.",
        [("Forums", [("table", ["Forum","Chair","Cadence","Mandate"],
            [["Group ExCo","Maria Donnelly","Monthly","Capital allocation & strategy"],
             ["Transformation SteerCo","Marcus Reilly","Bi-weekly","Portfolio gates & funding"],
             ["Treasury Council","Elena Vasquez","Monthly","Liquidity, banking, Kyriba"],
             ["Architecture Review Board","Liam O'Sullivan","Monthly","Tech standards & integration"],
             ["Risk & Audit Committee","David Chen","Quarterly","Controls, SOX, cyber"],
             ["AI Governance Board","Aisha Bello","Monthly","Use-case & model approval"]])])],
        meta=M(f, "Marcus Reilly", "Corporate Strategy"))
    reg(p, f, "Marcus Reilly", "Corporate Strategy", "Steering Committees", "Governance forums and mandates.")
    p = SRC / f / "raci_decision_rights.xlsx"
    decisions = ["Capital allocation","Bank selection","Kyriba scope change","ERP change request",
                 "Report retirement","Vendor renewal","AI use-case approval","Control exception",
                 "Hedging policy","Data product publish"]
    roles = ["CEO","CFO","CIO","CPO","Treasurer","CRO","CDO","SteerCo"]
    rows = [[d]+[random.choice(["R","A","C","I","-"]) for _ in roles] for d in decisions]
    write_xlsx(p, {"RACI": (["decision"]+roles, rows)}, title="RACI / Decision Rights", meta=M(f, "Marcus Reilly", "Corporate Strategy"))
    reg(p, f, "Marcus Reilly", "Corporate Strategy", "RACI / Decision Rights", "Decision-rights matrix across executive roles.")
    p = SRC / f / "org_design_narrative.docx"
    write_docx(p, "Organization Design Narrative",
        [("Current state", ["Finance and IT operate with overlapping demand-management processes; "
          "decision rights for cross-BU technology are ambiguous, slowing the Kyriba rollout and reporting changes."]),
         ("Target state", [("bullets",["Clear A (accountable) owner per enterprise decision type.",
            "Single intake for finance technology demand.","Federated data ownership with domain stewards."])])],
        meta=M(f, "Marcus Reilly", "Corporate Strategy"))
    reg(p, f, "Marcus Reilly", "Corporate Strategy", "Org Design Narrative", "Current vs target organization design.")

# ---- 03 strategy & initiatives --------------------------------------------
def build_strategy():
    f = "03_strategy_initiatives"
    p = SRC / f / "corporate_strategy_2026_2027.pdf"
    write_pdf(p, "Corporate Strategy 2026-2027", "Board-approved strategy on a page (synthetic).",
        [("Strategic pillars", [("bullets",[
            "Profitable growth in industrial and health-supplies segments.",
            "Cash discipline and treasury modernization.",
            "Digital & data backbone consolidation.",
            "Disciplined capital allocation and portfolio shaping."])]),
         ("Financial targets", [("table",["Metric","FY25","FY27 target"],
            [["EBITDA margin","11.8%","14.5%"],["Cash conversion","68%","82%"],
             ["IT run cost","$214M","$182M"],["Reports (managed)","320","140"]])])],
        meta=M(f, "Maria Donnelly", "Corporate Strategy"))
    reg(p, f, "Maria Donnelly", "Corporate Strategy", "Corporate Strategy 2026-2027", "Strategy pillars and financial targets.", sensitivity="Confidential")
    p = SRC / f / "board_priorities.pdf"
    write_pdf(p, "Board Priorities", "Board-level priorities and watch items.",
        [("Priorities",[("bullets",["Liquidity resilience","SOX & controls","Cyber posture",
            "AI governance","Cost out","Talent & succession"])])], meta=M(f,"Maria Donnelly","Board"))
    reg(p, f, "Maria Donnelly", "Board", "Board Priorities", "Board priorities and watch items.", sensitivity="Confidential")
    inits = ["Kyriba Treasury Modernization","SOX Controls Uplift","Reporting Rationalization",
             "SI/AMS Vendor Optimization","Snowflake Consolidation","S/4HANA Migration Wave 2",
             "Enterprise AI Platform","Order-to-Cash Automation","Procure-to-Pay on Coupa",
             "Cyber Zero-Trust","Data Quality Remediation","Close Acceleration"]
    p = SRC / f / "transformation_portfolio.xlsx"
    rows=[[f"INIT-{i+1:03d}", n, random.choice(["Treasury","Finance","IT","Procurement","Data","Risk"]),
           random.choice(["Charter","Discover","Design","Roadmap","Mobilize"]),
           random.choice(EXECS)[0], round(random.uniform(1.5,18),1), round(random.uniform(3,42),1),
           random.choice(["Green","Amber","Red"])] for i,n in enumerate(inits)]
    write_xlsx(p, {"Portfolio": (["init_id","initiative","theme","phase","sponsor","spend_m","benefit_m","rag"], rows)},
               title="Transformation Portfolio", meta=M(f,"Marcus Reilly","Corporate Strategy"))
    reg(p, f, "Marcus Reilly", "Corporate Strategy", "Transformation Portfolio", "Initiative portfolio with spend/benefit/RAG.", sensitivity="Confidential")
    p = SRC / f / "initiative_inventory.xlsx"
    rows=[[f"INIT-{i+1:03d}", n, date_in(), date_in("2026-07-01","2027-12-31"),
           random.choice(SI_PARTNERS), random.randint(4,40)] for i,n in enumerate(inits)]
    write_xlsx(p, {"Inventory": (["init_id","initiative","start","target_end","si_partner","fte"], rows)},
               title="Initiative Inventory", meta=M(f,"Marcus Reilly","Corporate Strategy"))
    reg(p, f, "Marcus Reilly", "Corporate Strategy", "Initiative Inventory", "Initiative timing, partners, staffing.")
    p = SRC / f / "budget_assumptions.xlsx"
    rows=[["Discount rate","WACC","9.5%"],["Benefit ramp","Months to 80%","14"],
          ["Contingency","% of spend","12%"],["FX","USD/EUR","1.08"],["Inflation","Annual","3.2%"]]
    write_xlsx(p, {"Assumptions": (["category","parameter","value"], rows)}, title="Budget Assumptions", meta=M(f,"Raymond Okafor","FP&A"))
    reg(p, f, "Raymond Okafor", "FP&A", "Budget Assumptions", "Planning assumptions for the portfolio.")
    p = SRC / f / "dependency_map.csv"
    rows=[[f"INIT-{i+1:03d}", f"INIT-{random.randint(1,len(inits)):03d}", random.choice(["blocks","enables","shares-resource"])]
          for i in range(len(inits)) for _ in range(random.randint(1,3))]
    write_csv(p, ["from_init","to_init","relationship"], rows, meta=M(f,"Marcus Reilly","Corporate Strategy"))
    reg(p, f, "Marcus Reilly", "Corporate Strategy", "Dependency Map", "Initiative dependency edges.")
    p = SRC / f / "raid_log.xlsx"
    rows=[[f"RAID-{i+1:03d}", random.choice(["Risk","Assumption","Issue","Dependency"]),
           random.choice(["Kyriba bank connectivity slippage","ERP freeze window conflict",
            "Data quality blocks reporting","SI ramp delay","Control owner capacity",
            "FX volatility","Cyber finding open","Cloud cost overrun"]),
           random.choice(["High","Medium","Low"]), random.choice(EXECS)[0], random.choice(["Open","Mitigating","Closed"])]
          for i in range(28)]
    write_xlsx(p, {"RAID": (["id","type","description","severity","owner","status"], rows)}, title="RAID Log", meta=M(f,"Marcus Reilly","Corporate Strategy"))
    reg(p, f, "Marcus Reilly", "Corporate Strategy", "RAID Log", "Risks, assumptions, issues, dependencies.")
    p = SRC / f / "strategy_narrative.docx"
    write_docx(p, "Strategy Narrative",
        [("Where we play",["Lakeshore concentrates capital on industrial components and health supplies "
          "while harvesting mature consumer and logistics lines."]),
         ("How we win",[("bullets",["Operational excellence","Cash & cost discipline","Data-led decisions","Governed AI"])])],
        meta=M(f,"Maria Donnelly","Corporate Strategy"))
    reg(p, f, "Maria Donnelly", "Corporate Strategy", "Strategy Narrative", "Where-we-play / how-we-win narrative.")
    # roadmap SVG built in rich_diagrams()

print("built: org, strategy")

# ---- 04 finance & performance ---------------------------------------------
def build_finance():
    f = "04_finance_performance"; owner="Raymond Okafor"
    p = SRC / f / "pnl_baseline.xlsx"
    lines=["Revenue","COGS","Gross Profit","SG&A","R&D","Other opex","EBIT","Interest","Tax","Net Income"]
    rows=[[l]+[round(random.uniform(40,8400),1) for _ in range(12)] for l in lines]
    write_xlsx(p, {"PnL_Monthly": (["line_item"]+[f"M{i+1:02d}" for i in range(12)], rows)},
               title="P&L Baseline (Monthly)", meta=M(f,owner,"Oracle Hyperion"))
    reg(p, f, owner, "Oracle Hyperion", "P&L Baseline", "Monthly group P&L baseline.", sensitivity="Confidential")
    p = SRC / f / "cash_working_capital.xlsx"
    rows=[[g[1], round(random.uniform(20,900),1), random.randint(38,72), random.randint(28,64), random.randint(40,95)]
          for g in GEOS]
    write_xlsx(p, {"WorkingCapital": (["country","cash_m","dso_days","dpo_days","dio_days"], rows)},
               title="Cash & Working Capital", meta=M(f,"Elena Vasquez","Kyriba TMS"))
    reg(p, f, "Elena Vasquez", "Kyriba TMS", "Cash & Working Capital", "Cash and working-capital metrics by country.", sensitivity="Confidential")
    for fn, title in [("fpa_process.docx","FP&A Process"),("forecasting_process.docx","Forecasting Process"),
                      ("close_process.docx","Financial Close Process")]:
        p = SRC / f / fn
        write_docx(p, title, [("Process",[("bullets",[
            "Inputs sourced from SAP and Hyperion; consolidated in OneStream.",
            "Cycle time and manual touchpoints are the primary pain points.",
            "Targeted automation via reporting rationalization and Blackline."])]),
            ("Pain points",[f"{title} suffers from spreadsheet sprawl, late actuals, and reconciliation rework."])],
            meta=M(f,"Sofia Marchetti","OneStream"))
        reg(p, f, "Sofia Marchetti", "OneStream", title, f"{title} narrative and pain points.")
    p = SRC / f / "management_reporting.pdf"
    write_pdf(p, "Management Reporting Overview", "Current reporting estate and pain points.",
        [("Estate",[("table",["Layer","Count","Tool"],
            [["Executive packs","18","Power BI"],["BU dashboards","96","Power BI"],
             ["Operational reports","206","SAP / Excel"]])]),
         ("Pain points",[("bullets",["Duplicated metrics","Manual data wrangling","No single semantic layer"])])],
        meta=M(f,"Sofia Marchetti","Power BI"))
    reg(p, f, "Sofia Marchetti", "Power BI", "Management Reporting", "Reporting estate overview.")
    p = SRC / f / "kpi_catalog.xlsx"
    kpis=["Revenue growth","EBITDA margin","Cash conversion","DSO","DPO","Inventory turns",
          "Forecast accuracy","Close cycle days","IT cost ratio","On-time payments","Hedge ratio","Audit findings"]
    rows=[[f"KPI-{i+1:03d}", k, random.choice(["Finance","Treasury","Ops","IT"]),
           random.choice(["Monthly","Quarterly"]), random.choice(EXECS)[0], f"{random.uniform(1,95):.1f}"] for i,k in enumerate(kpis)]
    write_xlsx(p, {"KPIs": (["kpi_id","kpi","domain","frequency","owner","current_value"], rows)}, title="KPI Catalog", meta=M(f,owner,"Power BI"))
    reg(p, f, owner, "Power BI", "KPI Catalog", "Enterprise KPI catalog.")
    p = SRC / f / "value_pools.xlsx"
    pools=["Idle cash reduction","Bank fee optimization","Reporting effort","AMS run cost",
           "Working capital","Procurement savings","Close acceleration","AI productivity"]
    rows=[[f"VP-{i+1:02d}", n, round(random.uniform(4,62),1), random.choice(["High","Medium"]),
           random.choice(["12-18m","6-12m","18-24m"])] for i,n in enumerate(pools)]
    write_xlsx(p, {"ValuePools": (["vp_id","value_pool","benefit_m","confidence","horizon"], rows)}, title="Value Pools", meta=M(f,owner,"FP&A"))
    reg(p, f, owner, "FP&A", "Value Pools", "Addressable value pools.", sensitivity="Confidential")
    p = SRC / f / "monthly_management_pack.pdf"
    write_pdf(p, "Monthly Management Pack (Illustrative)", "Synthetic management pack excerpt.",
        [("Group performance",[("table",["Metric","Actual","Plan","Var"],
            [["Revenue ($M)","712","728","-2.2%"],["EBITDA ($M)","98","104","-5.8%"],
             ["Cash ($M)","486","450","+8.0%"]])])], meta=M(f,owner,"Power BI"))
    reg(p, f, owner, "Power BI", "Monthly Management Pack", "Illustrative monthly pack.", sensitivity="Confidential")
    p = SRC / f / "finance_calendar.csv"
    rows=[[f"WD+{d}", random.choice(["Sub-ledger close","GL close","Consolidation","Reporting","Review"]),
           random.choice(EXECS)[0]] for d in range(1,11)]
    write_csv(p, ["close_day","activity","owner"], rows, meta=M(f,"Sofia Marchetti","OneStream"))
    reg(p, f, "Sofia Marchetti", "OneStream", "Finance Calendar", "Close calendar activities.")
    p = SRC / f / "working_capital_drivers.xlsx"
    rows=[[b[1], random.randint(40,72), random.randint(28,60), round(random.uniform(5,90),1)] for b in BUSINESS_UNITS]
    write_xlsx(p, {"WCDrivers": (["business_unit","dso","dpo","trapped_cash_m"], rows)}, title="Working Capital Drivers", meta=M(f,"Elena Vasquez","Kyriba TMS"))
    reg(p, f, "Elena Vasquez", "Kyriba TMS", "Working Capital Drivers", "WC drivers by BU.", sensitivity="Confidential")

# ---- 05 treasury & kyriba -------------------------------------------------
def build_treasury():
    f = "05_treasury_kyriba"; owner="Elena Vasquez"; sys="Kyriba TMS"
    p = SRC / f / "kyriba_rollout_plan.xlsx"
    waves=["Wave 1 NA","Wave 2 EMEA","Wave 3 APAC","Wave 4 LATAM"]
    rows=[[f"W{i+1}", w, date_in("2025-06-01","2026-03-01"), date_in("2026-03-01","2027-06-01"),
           random.choice(["Cash visibility","Payments","Liquidity","Bank connectivity"]),
           random.choice(["Green","Amber","Red"]), random.randint(4,18)] for i,w in enumerate(waves)]
    write_xlsx(p, {"Rollout": (["wave_id","wave","start","end","scope","rag","banks"], rows)}, title="Kyriba Rollout Plan", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Kyriba Rollout Plan", "Wave plan for Kyriba rollout.", sensitivity="Confidential")
    p = SRC / f / "bank_connectivity_inventory.xlsx"
    rows=[]
    for i,bnk in enumerate(BANKS):
        rows.append([f"BANK-{i+1:02d}", bnk, random.choice(["SWIFT","H2H","API","EBICS"]),
                     random.choice(["Live","In progress","Planned"]), random.randint(2,40),
                     random.choice(["MT940","MT942","CAMT.053","BAI2"])])
    write_xlsx(p, {"BankConnectivity": (["id","bank","channel","status","accounts","statement_format"], rows)},
               title="Bank Connectivity Inventory", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Bank Connectivity Inventory", "Bank channels and statement formats.", sensitivity="Confidential")
    p = SRC / f / "payment_process.docx"
    write_docx(p, "Payment Process & Payments Factory",
        [("End-to-end flow",["The Lakeshore payments factory centralizes outbound payments through Kyriba. AP proposals "
          "originate in SAP (IDoc PEXR2002), are enriched and validated against entity-level limits in Kyriba, pass dual "
          "authorization and real-time sanctions screening, and are transmitted to banks as ISO20022 pain.001 over SFTP/SWIFT. "
          "Bank acknowledgements (pain.002) and statements (camt.053/052) flow back for auto-reconciliation to the GL."]),
         ("Process steps",[("table",["Step","System","Control","SLA"],
            [["AP proposal","SAP","3-way match","T+0"],["Enrichment/limits","Kyriba","limit check","<5m"],
             ["Dual authorization","Kyriba","SoD","same day"],["Sanctions screening","Kyriba/AML","OFAC/EU","real-time"],
             ["Transmit pain.001","Kyriba→Bank","format validation","intraday"],["Recon camt.053","Bank→Kyriba→SAP","auto-match","T+1"]])]),
         ("Volumetrics & insight",["~412k payments/year across 10 banks; 6.8% exception rate (target <3%). Primary exception "
          "drivers: beneficiary bank detail mismatches and missing remittance refs — addressed by vendor master cleanup and "
          "payment anomaly detection (see AI use cases)."])],
        meta=M(f,owner,"Kyriba TMS"))
    reg(p, f, owner, "Kyriba TMS", "Payment Process", "Payments factory flow, controls, volumetrics.", sensitivity="Confidential")
    p = SRC / f / "kyriba_integration_design.docx"
    write_docx(p, "Kyriba Integration Design",
        [("Interfaces",[("table",["Interface","Direction","Format","Channel","Frequency"],
            [["Payment proposals","SAP→Kyriba","IDoc PEXR2002","ESB","hourly"],
             ["Outbound payments","Kyriba→Bank","pain.001.001.09","SFTP/SWIFT","intraday"],
             ["Payment status","Bank→Kyriba","pain.002","SFTP","intraday"],
             ["Statements","Bank→Kyriba","camt.053/052","SFTP","daily/intraday"],
             ["GL recon postings","Kyriba→SAP","BAPI/IDoc","ESB","T+1"],
             ["FX rates","Refinitiv→Kyriba","REST","API","intraday"]])]),
         ("Connectivity",["Kyriba bank connectivity uses SWIFT Alliance Lite2 (SCORE) for SWIFT-enabled banks and host-to-host "
          "SFTP/EBICS for the remainder; Kyriba's format library handles country-specific ISO20022 variants. Mutual TLS and "
          "PGP protect file channels."]),
         ("Design principles",[("bullets",["Single payment hub; no direct ERP-to-bank channels.",
            "Idempotent, replayable interfaces with full audit trail.","Clean-core: transformations in Kyriba/ESB, not ERP.",
            "Environment parity (DEV/QA/PROD) with masked test data."])])],
        meta=M(f,owner,"SAP S/4HANA"))
    reg(p, f, owner, "SAP S/4HANA", "Kyriba Integration Design", "Kyriba interface and connectivity design.", sensitivity="Confidential")
    p = SRC / f / "treasury_policy.docx"
    write_docx(p, "Treasury Policy",
        [("Scope",["Governs cash management, banking, payments, FX/interest-rate risk, and counterparty limits for all "
          "Lakeshore entities."]),
         ("Key policies",[("bullets",["Minimum cash buffers by entity and currency.",
            "Hedging bands for transactional FX exposure (50-90% coverage).","Counterparty concentration limits per bank.",
            "Dual authorization and four-eyes on all outbound payments.","Quarterly bank account mandate review."])])],
        meta=M(f,owner,"Treasury"))
    reg(p, f, owner, "Treasury", "Treasury Policy", "Treasury policy framework.", sensitivity="Confidential")
    p = SRC / f / "cash_positioning.xlsx"
    rows=[[g[1], round(random.uniform(5,400),1), round(random.uniform(0,120),1), random.choice(BANKS)] for g in GEOS]
    write_xlsx(p, {"CashPosition": (["country","cash_m","idle_cash_m","concentration_bank"], rows)}, title="Cash Positioning", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Cash Positioning", "Cash position by country.", sensitivity="Confidential")
    p = SRC / f / "liquidity_forecasting.xlsx"
    rows=[[f"W{w}", round(random.uniform(300,560),1), round(random.uniform(-60,80),1), f"{random.uniform(70,96):.1f}%"] for w in range(1,14)]
    write_xlsx(p, {"Liquidity": (["week","forecast_cash_m","net_flow_m","forecast_accuracy"], rows)}, title="Liquidity Forecasting", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Liquidity Forecasting", "13-week liquidity forecast.", sensitivity="Confidential")
    p = SRC / f / "debt_fx_exposure.xlsx"
    rows=[[c, round(random.uniform(20,600),1), f"{random.uniform(2,7):.2f}%", random.choice(["Fixed","Floating"])]
          for c in ["USD","EUR","GBP","SGD","INR","BRL"]]
    write_xlsx(p, {"DebtFX": (["currency","exposure_m","rate","type"], rows)}, title="Debt & FX Exposure", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Debt & FX Exposure", "Debt and FX exposures.", sensitivity="Confidential")
    p = SRC / f / "bank_fee_baseline.xlsx"
    rows=[[bnk, round(random.uniform(0.1,3.4),2), random.randint(200,9000), round(random.uniform(0.05,0.9),2)] for bnk in BANKS]
    write_xlsx(p, {"BankFees": (["bank","annual_fee_m","transactions_k","avg_fee_usd"], rows)}, title="Bank Fee Baseline", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Bank Fee Baseline", "Annual bank fees baseline.", sensitivity="Confidential")
    p = SRC / f / "treasury_controls.pdf"
    write_pdf(p, "Treasury Controls Summary", "Key treasury controls (synthetic).",
        [("Controls",[("table",["Control","Type","Status"],
            [["Payment dual auth","Preventive","Effective"],["Sanctions screening","Detective","Effective"],
             ["Bank account mandate review","Preventive","Gap"],["Cash forecast review","Detective","Partial"]])])],
        meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Treasury Controls", "Treasury control summary.", sensitivity="Confidential")
    p = SRC / f / "treasury_target_operating_model.pdf"
    write_pdf(p, "Treasury Target Operating Model", "Future-state treasury TOM.",
        [("Pillars",[("bullets",["Centralized cash visibility","In-house bank","Automated payments factory","Real-time liquidity"])])],
        meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Treasury TOM", "Treasury target operating model.")
    p = SRC / f / "kyriba_rollout_risks_defects.xlsx"
    rows=[[f"KRISK-{i+1:03d}", random.choice(["Connectivity","Data mapping","Testing","Change mgmt","Security"]),
           random.choice(["Open","Mitigating","Closed"]), random.choice(["High","Medium","Low"]),
           random.choice(BANKS)] for i in range(24)]
    write_xlsx(p, {"RolloutRisks": (["id","category","status","severity","related_bank"], rows)}, title="Kyriba Rollout Risks/Defects", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Kyriba Rollout Risks/Defects", "Rollout risks and defects.", sensitivity="Confidential")
    p = SRC / f / "kyriba_test_defects.csv"
    rows=[[f"DEF-{i+1:04d}", random.choice(["P1","P2","P3","P4"]),
           random.choice(["Statement import fails","Payment file rejected","FX rate mismatch","Reconciliation gap","Timeout"]),
           random.choice(["Open","Fixed","Retest","Closed"]), random.choice(BANKS), date_in("2025-09-01","2026-06-01")]
          for i in range(60)]
    write_csv(p, ["defect_id","priority","summary","status","bank","raised_date"], rows, meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Kyriba Test Defects", "Test defect log for Kyriba.", sensitivity="Confidential")

print("built: finance, treasury")

# ---- 06 IT systems & architecture -----------------------------------------
def build_it():
    f = "06_it_systems_architecture"; owner="Liam O'Sullivan"
    p = SRC / f / "application_inventory.xlsx"
    apps=[]
    for i in range(120):
        apps.append([f"APP-{i+1:04d}",
            random.choice(["SAP ECC","S/4HANA","Workday","Coupa","Kyriba","ServiceNow","Snowflake",
                "Power BI","Salesforce","Hyperion","Blackline","Legacy AS/400","Custom .NET","Tableau","HighRadius"])
            + f" {random.choice(['Core','Module','Instance'])} {i%9}",
            random.choice(["Finance","Treasury","HR","Procurement","Sales","IT","Data","Ops"]),
            random.choice(["Strategic","Tolerate","Migrate","Retire"]),
            random.choice(["SaaS","On-prem","IaaS","PaaS"]),
            random.choice(EXECS)[0], round(random.uniform(0.05,4.2),2),
            random.choice(["Low","Medium","High","Critical"])])
    write_xlsx(p, {"AppInventory": (["app_id","application","function","disposition","hosting","owner","annual_cost_m","criticality"], apps)},
               title="Application Inventory", meta=M(f,owner,"ServiceNow CMDB"))
    reg(p, f, owner, "ServiceNow CMDB", "Application Inventory", "120-row application portfolio.")
    p = SRC / f / "erp_landscape.pdf"
    write_pdf(p, "ERP Landscape & S/4HANA Migration", "SAP estate, technical debt, and brownfield migration approach.",
        [("Production instances", [
            "Lakeshore operates a federated SAP estate: two ECC 6.0 instances (NA, EMEA) on Oracle DB, one S/4HANA 2022 "
            "instance (APAC) on HANA, and a legacy ECC instance (LATAM). Consolidation to a single S/4HANA tenant is the "
            "FY26-27 ERP north star.",
            ("table",["Instance","Region","Version","DB","Modules","Migration","Interfaces"],
            [["PRD-NA","North America","ECC 6.0","Oracle 19c","FI/CO/MM/SD/PP","Wave 2 FY26 (brownfield)","98"],
             ["PRD-EU","EMEA","ECC 6.0","Oracle 19c","FI/CO/MM/SD","Wave 2 FY26 (brownfield)","74"],
             ["PRD-AP","APAC","S/4HANA 2022","HANA 2.0","Finance/Sourcing","Live (selective data transition)","41"],
             ["PRD-LA","LATAM","ECC 6.0","Oracle 12c","FI/MM","Wave 3 FY27","27"]])]),
         ("Integration & data flows", [
            "ERP integrates outbound via IDoc (PEXR2002 payments to Kyriba, INVOIC, ORDERS) and OData/BAPI for real-time "
            "lookups; inbound master data via ALE distribution. Financial actuals flow nightly to Snowflake via IDoc→Parquet "
            "landing, then dbt curation. Payment proposals flow hourly to Kyriba over the ESB.",
            ("bullets",["Outbound: IDoc (payments, invoices), OData (real-time), ALE (master data)",
                        "Inbound: bank statements via Kyriba→ERP recon postings",
                        "Batch: GL/AP/AR nightly extract → Snowflake bronze",
                        "Clean-core principle for S/4: extensions on BTP, no core mods"])]),
         ("Technical debt & risks", [("bullets",[
            "ECC mainstream maintenance end-of-life pressure; 1,400+ custom Z-objects to remediate.",
            "Inconsistent chart-of-accounts across instances complicates consolidation.",
            "Interface sprawl (240+) raises migration regression risk — see integration architecture."])]),
        ], meta=M(f,"Priya Natarajan","SAP"))
    reg(p, f, "Priya Natarajan", "SAP", "ERP Landscape", "ERP instances, integration/data flows, migration approach, tech debt.", sensitivity="Confidential")
    for fn,title,system in [("hris_workday.docx","HRIS / Workday","Workday HCM"),
                            ("procurement_systems.docx","Procurement Systems","Coupa"),
                            ("identity_sso.docx","Identity & SSO","Azure Active Directory")]:
        p = SRC / f / fn
        write_docx(p, title, [("Overview",[f"{title} landscape for Lakeshore (synthetic)."]),
            ("Notes",[("bullets",["Integration via middleware (ESB/APIM)","Identity federated to Entra ID (SAML/OIDC)","Data lands in Snowflake via CDC/EIB"])])],
            meta=M(f,owner,system))
        reg(p, f, owner, system, title, f"{title} overview.")
    p = SRC / f / "data_platforms.docx"
    write_docx(p, "Data Platforms",
        [("Architecture",["Lakeshore's analytics platform is a Snowflake medallion (bronze/silver/gold) fed from ADLS gen2 "
          "landing, transformed with dbt, governed with RBAC + dynamic masking, and consumed by Power BI, Azure AI Search "
          "(tenant-context-v1), and the Postgres context layer. See `data_platform_architecture.svg`."]),
         ("Components",[("table",["Layer","Technology","Purpose"],
            [["Landing/Bronze","ADLS gen2 + Snowpipe","immutable raw ingest"],
             ["Curated/Silver","Snowflake + dbt","conformed, tested models"],
             ["Present/Gold","Snowflake marts","domain marts + exposures"],
             ["Semantic","dbt metrics","certified KPI definitions"],
             ["Serving","Power BI / AI Search / Postgres","BI + RAG + context"]])]),
         ("Insight",["~58% of report effort is manual data prep upstream of certified marts; consolidating onto the semantic "
          "layer is the single largest reporting-rationalization lever."])],
        meta=M(f,"Aisha Bello","Snowflake"))
    reg(p, f, "Aisha Bello", "Snowflake", "Data Platforms", "Medallion data platform components and insight.")
    p = SRC / f / "finance_systems.xlsx"
    rows=[["SAP FI","GL/AP/AR","SAP ECC","Migrate"],["Hyperion","Consolidation","Oracle","Tolerate"],
          ["OneStream","FP&A","OneStream","Strategic"],["Blackline","Reconciliation","Blackline","Strategic"],
          ["HighRadius","Cash app","HighRadius","Strategic"],["Kyriba","Treasury","Kyriba","Strategic"]]
    write_xlsx(p, {"FinanceSystems": (["system","capability","vendor","disposition"], rows)}, title="Finance Systems", meta=M(f,"Raymond Okafor","SAP"))
    reg(p, f, "Raymond Okafor", "SAP", "Finance Systems", "Finance system map.")
    p = SRC / f / "reporting_tools.xlsx"
    rows=[["Power BI","BI","2400","Strategic"],["Tableau","BI","320","Migrate"],["SAP BW","DW","-","Retire"],
          ["Excel","Adhoc","-","Tolerate"],["Snowflake","Platform","-","Strategic"]]
    write_xlsx(p, {"ReportingTools": (["tool","category","users","disposition"], rows)}, title="Reporting Tools", meta=M(f,"Aisha Bello","Power BI"))
    reg(p, f, "Aisha Bello", "Power BI", "Reporting Tools", "Reporting tool inventory.")
    p = SRC / f / "security_tools.xlsx"
    rows=[["Microsoft Defender","EDR"],["Sentinel","SIEM"],["SailPoint","IGA"],["CyberArk","PAM"],
          ["Zscaler","SSE"],["Qualys","VM"],["Proofpoint","Email"]]
    rows=[[t,c,random.choice(["Deployed","Partial","Planned"])] for t,c in rows]
    write_xlsx(p, {"SecurityTools": (["tool","category","status"], rows)}, title="Security Tools", meta=M(f,"David Chen","Sentinel"))
    reg(p, f, "David Chen", "Sentinel", "Security Tools", "Security tool inventory.")
    p = SRC / f / "legacy_applications.xlsx"
    rows=[[f"LEG-{i+1:03d}", random.choice(["AS/400 Billing","VB6 Pricing","Access DB","Mainframe GL","Lotus Notes"])+f" {i}",
           random.randint(8,28), random.choice(["High","Medium","Critical"]), "Retire"] for i in range(20)]
    write_xlsx(p, {"Legacy": (["id","application","age_years","risk","disposition"], rows)}, title="Legacy Applications", meta=M(f,owner,"ServiceNow CMDB"))
    reg(p, f, owner, "ServiceNow CMDB", "Legacy Applications", "Legacy application risk inventory.")
    p = SRC / f / "integration_architecture.pdf"
    write_pdf(p, "Integration Architecture Narrative", "Backbone topology, patterns, standards, and the interface estate.",
        [("Topology", [
            "Lakeshore runs a hybrid integration backbone: SAP PI/PO (ESB) for ERP-centric A2A/B2B, Azure API Management "
            "for REST facades to SaaS, Azure Service Bus/Event Grid for event-driven ingestion (including the "
            "q-context-ingestion-events queue used by the context loader), and a Managed File Transfer tier for bank "
            "host-to-host. See `integration_architecture_diagram.svg` for the component view and interface inventory."]),
         ("Patterns & standards", [("table",["Pattern","Where used","Standard/Protocol"],
            [["Request/response","SaaS lookups via APIM","REST/JSON, OAuth2, mTLS"],
             ["Async messaging","Context ingestion, eventing","AMQP (Service Bus), pub/sub, DLQ"],
             ["File-based batch","Bank + ERP bulk","SFTP/MFT, PGP, ISO20022 XML"],
             ["A2A orchestration","ERP↔TMS↔EPM","IDoc, SOAP, XSLT mapping"],
             ["CDC/ELT","Source→Snowflake","Snowpipe, log-based CDC"]])]),
         ("Key end-to-end flows", [("bullets",[
            "Payments: SAP IDoc → Kyriba → ISO20022 pain.001 → bank → pain.002 ack → camt.053 statement → ERP recon.",
            "Context load: ZIP → Blob (context-drops) → Service Bus → worker → parse → enterprise_context_chunks → AI Search.",
            "Analytics: SAP/Workday/Coupa → Snowflake medallion → semantic layer → Power BI + AI Search."])]),
         ("Resilience & observability", [
            "Guaranteed-once delivery via Service Bus sessions + idempotent sinks; dead-letter queues with replay; "
            "interface SLAs monitored in Azure Monitor with incidents routed to ServiceNow. 240+ active interfaces are "
            "inventoried with owner, criticality, and recovery objectives."]),
        ], meta=M(f,owner,"SAP PI/PO"))
    reg(p, f, owner, "SAP PI/PO", "Integration Architecture", "Backbone topology, patterns, standards, end-to-end flows.")
    # interface inventory (structured)
    p = SRC / f / "interface_inventory.xlsx"
    irows=[]
    systems_pairs=[("SAP ECC","Kyriba"),("SAP ECC","Snowflake"),("Kyriba","JPMorgan"),("Coupa","SAP ECC"),
                   ("Workday","SAP ECC"),("ServiceNow","Snowflake"),("Salesforce","SAP ECC"),("Snowflake","Power BI"),
                   ("Snowflake","Azure AI Search"),("Kyriba","Citi")]
    fmts=["IDoc PEXR2002","ISO20022 pain.001","ISO20022 camt.053","cXML","REST/JSON","OData","Bulk API","Parquet","JDBC upsert","SFTP/PGP"]
    for i in range(64):
        a,b=random.choice(systems_pairs)
        irows.append([f"IF-{i+1:04d}", a, b, random.choice(fmts),
                      random.choice(["real-time","event","batch","micro-batch"]),
                      random.choice(["ESB","APIM","Service Bus","MFT","CDC"]),
                      random.choice(["Critical","High","Medium","Low"]),
                      random.choice(["1m","15m","1h","4h","24h"]), random.choice(EXECS)[0]])
    write_xlsx(p, {"Interfaces": (["interface_id","source","target","format_standard","pattern","channel","criticality","rpo","owner"], irows)},
               title="Interface Inventory", meta=M(f,owner,"SAP PI/PO"))
    reg(p, f, owner, "SAP PI/PO", "Interface Inventory", "64-row integration interface estate with protocols and RPO.")
    # current-state architecture SVG built in rich_diagrams()

# ---- 07 data, analytics, reporting ----------------------------------------
def build_data():
    f = "07_data_analytics_reporting"; owner="Aisha Bello"
    p = SRC / f / "report_inventory.xlsx"
    rows=[]
    for i in range(320):
        rows.append([f"RPT-{i+1:04d}",
            random.choice(["Exec","BU","Operational","Regulatory","Adhoc"])+f" report {i}",
            random.choice(["Finance","Treasury","Sales","Ops","HR","IT"]),
            random.choice(["Power BI","Tableau","SAP BW","Excel"]),
            random.choice(["Daily","Weekly","Monthly","Quarterly"]),
            random.randint(2,2200), random.choice(["Keep","Consolidate","Retire","Rebuild"]),
            random.choice(EXECS)[0]])
    write_xlsx(p, {"Reports": (["report_id","report_name","domain","tool","frequency","consumers","disposition","owner"], rows)},
               title="Report Inventory", meta=M(f,owner,"Power BI"))
    reg(p, f, owner, "Power BI", "Report Inventory", "320-row report estate for rationalization.")
    p = SRC / f / "dashboard_inventory.xlsx"
    rows=[[f"DASH-{i+1:03d}", f"Dashboard {i}", random.choice(["Power BI","Tableau"]),
           random.randint(5,900), random.choice(["Active","Stale","Duplicate"])] for i in range(80)]
    write_xlsx(p, {"Dashboards": (["dash_id","name","tool","viewers","status"], rows)}, title="Dashboard Inventory", meta=M(f,owner,"Power BI"))
    reg(p, f, owner, "Power BI", "Dashboard Inventory", "Dashboard inventory.")
    p = SRC / f / "kpi_definitions.xlsx"
    rows=[[f"KPI-{i+1:03d}", k, f"Definition of {k.lower()} (synthetic).", random.choice(EXECS)[0], random.choice(["Certified","Draft","Conflicting"])]
          for i,k in enumerate(["Revenue","EBITDA","DSO","DPO","Cash conversion","Forecast accuracy","Close days","On-time pay","Hedge ratio","NPS"])]
    write_xlsx(p, {"KPIDefs": (["kpi_id","kpi","definition","steward","status"], rows)}, title="KPI Definitions", meta=M(f,owner,"Snowflake"))
    reg(p, f, owner, "Snowflake", "KPI Definitions", "Certified/conflicting KPI definitions.")
    p = SRC / f / "data_ownership.xlsx"
    rows=[[d, random.choice(EXECS)[0], random.choice(["Snowflake","SAP","Workday","Coupa"]), random.choice(["Defined","Gap"])]
          for d in ["Customer","Vendor","Product","GL","Cost center","Employee","Bank account","Contract"]]
    write_xlsx(p, {"DataOwnership": (["data_domain","steward","system_of_record","governance_status"], rows)}, title="Data Ownership", meta=M(f,owner,"Snowflake"))
    reg(p, f, owner, "Snowflake", "Data Ownership", "Data domain ownership.")
    p = SRC / f / "data_quality_issues.xlsx"
    rows=[[f"DQ-{i+1:03d}", random.choice(["Customer","Vendor","GL","Bank account"]),
           random.choice(["Duplicates","Missing values","Stale","Inconsistent mapping"]),
           random.choice(["High","Medium","Low"]), random.choice(["Open","Remediating","Closed"])] for i in range(40)]
    write_xlsx(p, {"DQIssues": (["id","domain","issue_type","severity","status"], rows)}, title="Data Quality Issues", meta=M(f,owner,"Snowflake"))
    reg(p, f, owner, "Snowflake", "Data Quality Issues", "Data quality issue register.")
    p = SRC / f / "lineage_map.csv"
    rows=[[random.choice(["SAP","Workday","Coupa","Kyriba"]), random.choice(["Snowflake RAW","Snowflake CURATED"]),
           random.choice(["Power BI","Tableau"]), f"flow_{i}"] for i in range(60)]
    write_csv(p, ["source","transform_layer","consumer","flow_id"], rows, meta=M(f,owner,"Snowflake"))
    reg(p, f, owner, "Snowflake", "Lineage Map", "Source-to-consumer lineage edges.")
    p = SRC / f / "semantic_layer_assumptions.docx"
    write_docx(p, "Semantic Layer Assumptions & Design",
        [("Design",["A dbt-based semantic layer is the single source of certified metrics for both BI and AI retrieval. "
          "Conformed dimensions (entity, cost center, GL account, vendor, currency, calendar) and certified metrics "
          "(revenue, EBITDA, DSO/DPO, cash conversion, forecast accuracy) are defined once and consumed everywhere."]),
         ("Assumptions",[("bullets",[
            "Single conformed dimension set across finance, treasury, and procurement marts.",
            "Certified metric definitions versioned in dbt with tests, ownership, and lineage exposures.",
            "BI tools and the AI context broker consume governed metrics only (no ad-hoc SQL in reports).",
            "Row/column access policies enforce tenant and sensitivity scoping at the warehouse.",
            "Metric changes flow through a governance PR with steward approval."])]),
         ("Conflicts to resolve",[("table",["Metric","Issue","Resolution"],
            [["Revenue","3 definitions across BUs","certify single definition"],
             ["Cash conversion","timing basis differs","standardize on trailing-90"],
             ["Forecast accuracy","numerator varies","define MAPE basis"]])])],
        meta=M(f,owner,"Snowflake"))
    reg(p, f, owner, "Snowflake", "Semantic Layer Assumptions", "Semantic layer design, assumptions, conflict resolution.")
    p = SRC / f / "data_products.xlsx"
    rows=[[f"DP-{i+1:02d}", n, random.choice(EXECS)[0], random.choice(["Live","Build","Concept"])]
          for i,n in enumerate(["Cash 360","Vendor 360","Customer 360","Close cockpit","Spend analytics","Treasury risk","Reporting mart"])]
    write_xlsx(p, {"DataProducts": (["dp_id","data_product","owner","status"], rows)}, title="Data Products", meta=M(f,owner,"Snowflake"))
    reg(p, f, owner, "Snowflake", "Data Products", "Data product catalog.")
    p = SRC / f / "analytics_backlog.xlsx"
    rows=[[f"AN-{i+1:03d}", random.choice(["Forecast model","Anomaly detection","Spend cube","Cash predictor","Churn model"])+f" {i}",
           random.choice(["High","Medium","Low"]), random.choice(["Backlog","In progress","Done"])] for i in range(30)]
    write_xlsx(p, {"Backlog": (["id","item","priority","status"], rows)}, title="Analytics Backlog", meta=M(f,owner,"Snowflake"))
    reg(p, f, owner, "Snowflake", "Analytics Backlog", "Analytics demand backlog.")
    p = SRC / f / "reporting_rationalization_brief.pdf"
    write_pdf(p, "Reporting Rationalization Brief", "Case to reduce 320 reports to ~140.",
        [("Opportunity",[("bullets",["46% of reports duplicate metrics","58% manual data prep","Target 55% effort reduction"])]),
         ("Approach",[("table",["Step","Outcome"],[["Inventory & usage","Fact base"],["Rationalize","Retire/consolidate"],
            ["Rebuild on semantic layer","Certified metrics"],["Automate","Reduced manual effort"]])])],
        meta=M(f,owner,"Power BI"))
    reg(p, f, owner, "Power BI", "Reporting Rationalization Brief", "Reporting rationalization business case.")

print("built: it, data")

# ---- 08 operations & business process --------------------------------------
def build_ops():
    f = "08_operations_business_process"; owner="Sofia Marchetti"
    procs=[("order_to_cash.docx","Order-to-Cash","Salesforce"),("procure_to_pay.docx","Procure-to-Pay","Coupa"),
           ("record_to_report.docx","Record-to-Report","SAP S/4HANA"),("treasury_operations.docx","Treasury Operations","Kyriba TMS"),
           ("vendor_onboarding.docx","Vendor Onboarding","Coupa"),("planning_cycle.docx","Planning Cycle","OneStream"),
           ("management_reporting_process.docx","Management Reporting Process","Power BI")]
    for fn,title,system in procs:
        p = SRC / f / fn
        write_docx(p, title, [("Process steps",[("bullets",[
            f"Trigger and intake for {title.lower()}.","Hand-offs across shared services and BUs.",
            "Controls and approvals embedded.","Known bottlenecks and automation candidates."])]),
            ("KPIs",[("table",["KPI","Current","Target"],
                [["Cycle time","high","-30%"],["Manual touches","high","-50%"],["Exception rate","medium","-25%"]])])],
            meta=M(f,owner,system))
        reg(p, f, owner, system, title, f"{title} process narrative and KPIs.")
    p = SRC / f / "month_end_close.pdf"
    write_pdf(p, "Month-End Close", "Close process and acceleration plan.",
        [("Current",[("bullets",["8.5 day close","Manual reconciliations","Late actuals from BUs"])]),
         ("Target",[("bullets",["5 day close","Blackline automation","Continuous accounting"])])], meta=M(f,owner,"OneStream"))
    reg(p, f, owner, "OneStream", "Month-End Close", "Close acceleration plan.")
    p = SRC / f / "process_kpi_baseline.xlsx"
    rows=[[n, random.choice(["O2C","P2P","R2R","Treasury"]), round(random.uniform(1,28),1), f"{random.uniform(60,99):.1f}%"]
          for n in ["Cycle time","First-pass yield","Exception rate","On-time","Rework rate","Automation rate"]]
    write_xlsx(p, {"ProcessKPIs": (["kpi","process","value","target_attainment"], rows)}, title="Process KPI Baseline", meta=M(f,owner,"ServiceNow"))
    reg(p, f, owner, "ServiceNow", "Process KPI Baseline", "Process KPI baseline.")

# ---- 09 servicenow / support / workload (HIGH VOLUME) ----------------------
def build_servicenow():
    f = "09_servicenow_support_workload"; owner="Priya Natarajan"; sys="ServiceNow ITSM"
    cats=["Incident-ERP","Incident-Treasury","Incident-Reporting","Incident-Network","Incident-Access",
          "Request-Access","Request-Report","Request-Change","Problem","Enhancement"]
    apps=["SAP ECC","S/4HANA","Kyriba","Power BI","Workday","Coupa","Snowflake","Azure AD","ServiceNow","Salesforce"]
    causes=["Config error","Data quality","Integration failure","User error","Capacity","Code defect",
            "Vendor outage","Permission","Network","Unknown"]
    # 1600 incidents
    p = SRC / f / "servicenow_incidents.csv"
    rows=[]
    for i in range(1600):
        opened=date_in("2025-01-01","2026-06-01")
        pr=random.choices(["P1","P2","P3","P4"], weights=[4,16,50,30])[0]
        rows.append([f"INC{100000+i}", random.choice(cats), random.choice(apps), pr,
            random.choices(["Resolved","Closed","Open","In Progress","On Hold"], weights=[45,35,8,8,4])[0],
            opened, round(random.uniform(0.2, 220), 1), random.choice(causes),
            random.choice(["L1","L2","L3"]), random.choice([g[2] for g in GEOS]),
            random.choice(EXECS)[0], f"{pr} {random.choice(['outage','error','slow','failure','request'])} in {random.choice(apps)}"])
    write_csv(p, ["number","category","application","priority","state","opened_at","resolve_hours",
                  "root_cause","tier","region","assignment_owner","short_description"], rows, meta=M(f,owner,sys))
    reg(p, f, owner, sys, "ServiceNow Incidents", "1,600 synthetic incidents for retrieval volume.")
    # JSONL events (high volume)
    p = SRC / f / "servicenow_events.jsonl"
    recs=[]
    for i in range(1200):
        recs.append({"event_id": f"EVT{200000+i}", "type": random.choice(["state_change","comment","reassign","sla_breach"]),
                     "number": f"INC{100000+random.randint(0,1599)}", "ts": date_in("2025-01-01","2026-06-01"),
                     "actor": random.choice(EXECS)[0], "detail": random.choice(["escalated","resolved","awaiting user","vendor engaged"])})
    write_jsonl(p, recs, meta=M(f,owner,sys))
    reg(p, f, owner, sys, "ServiceNow Events", "1,200 synthetic ITSM events (JSONL).")
    p = SRC / f / "enhancement_backlog.xlsx"
    rows=[[f"ENH-{i+1:03d}", random.choice(apps), f"Enhancement {i}", random.choice(["High","Medium","Low"]),
           random.randint(1,40), random.choice(["Backlog","Scoped","In Dev","Done"])] for i in range(120)]
    write_xlsx(p, {"Backlog": (["id","application","title","priority","story_points","status"], rows)}, title="Enhancement Backlog", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Enhancement Backlog", "120-row enhancement backlog.")
    p = SRC / f / "defects.xlsx"
    rows=[[f"BUG-{i+1:04d}", random.choice(apps), random.choice(["P1","P2","P3"]),
           random.choice(["Open","Fixed","Closed"]), random.choice(causes)] for i in range(90)]
    write_xlsx(p, {"Defects": (["id","application","severity","status","root_cause"], rows)}, title="Defects", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Defects", "Defect register.")
    p = SRC / f / "sla_model.xlsx"
    rows=[["P1","15m","4h","99.5%"],["P2","30m","8h","98%"],["P3","4h","3d","95%"],["P4","1d","5d","90%"]]
    write_xlsx(p, {"SLA": (["priority","response","resolution","target_attainment"], rows)}, title="SLA Model", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "SLA Model", "Support SLA model.")
    p = SRC / f / "ticket_categories.xlsx"
    rows=[[c, random.randint(40,900), f"{random.uniform(60,99):.1f}%"] for c in cats]
    write_xlsx(p, {"Categories": (["category","volume","sla_attainment"], rows)}, title="Ticket Categories", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Ticket Categories", "Ticket category volumes.")
    p = SRC / f / "ticket_aging.xlsx"
    rows=[[b, random.randint(10,400)] for b in ["0-1d","1-3d","3-7d","7-14d","14-30d","30d+"]]
    write_xlsx(p, {"Aging": (["age_bucket","open_count"], rows)}, title="Ticket Aging", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Ticket Aging", "Open ticket aging.")
    p = SRC / f / "root_cause_analysis.pdf"
    write_pdf(p, "Root Cause Analysis", "Top recurring root causes.",
        [("Top causes",[("table",["Root cause","Share","Action"],
            [["Integration failure","22%","Harden interfaces"],["Data quality","19%","DQ remediation"],
             ["Config error","17%","Change control"],["Capacity","11%","Right-size"]])])], meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Root Cause Analysis", "RCA of recurring incidents.")
    p = SRC / f / "support_team_structure.xlsx"
    rows=[[t, random.randint(8,60), random.choice(SI_PARTNERS), random.choice([g[2] for g in GEOS])] for t in ["L1 Service Desk","L2 App Support","L3 Engineering","Treasury Support","Reporting Support"]]
    write_xlsx(p, {"SupportTeams": (["team","fte","provider","location"], rows)}, title="Support Team Structure", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Support Team Structure", "Support org structure.")
    p = SRC / f / "automation_opportunities.xlsx"
    rows=[[f"AUTO-{i+1:02d}", n, round(random.uniform(0.2,3.5),1), random.choice(["High","Medium"])]
          for i,n in enumerate(["Password reset","Access provisioning","Report distribution","Ticket triage",
            "Statement reconciliation","Payment status","Knowledge deflection","Incident summarization"])]
    write_xlsx(p, {"Automation": (["id","opportunity","annual_savings_m","feasibility"], rows)}, title="Automation Opportunities", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Automation Opportunities", "Support automation opportunities.")

print("built: ops, servicenow")

# ---- 10 vendors, contracts, sourcing --------------------------------------
def build_vendors():
    f = "10_vendors_contracts_source"; owner="Tomas Halvorsen"; sys="Coupa"
    p = SRC / f / "si_contracts.xlsx"
    rows=[[f"SI-{i+1:03d}", random.choice(SI_PARTNERS), random.choice(["Kyriba","S/4HANA","Reporting","Data"]),
           round(random.uniform(2,40),1), date_in("2024-01-01","2025-06-01"), date_in("2026-06-01","2028-12-31"),
           random.choice(["T&M","Fixed","Outcome"])] for i in range(20)]
    write_xlsx(p, {"SIContracts": (["id","partner","program","tcv_m","start","end","commercial_model"], rows)}, title="SI Contracts", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "SI Contracts", "System integrator contracts.", sensitivity="Confidential")
    p = SRC / f / "software_contracts.xlsx"
    rows=[[f"SW-{i+1:03d}", random.choice(["SAP","Oracle","Microsoft","Salesforce","Snowflake","ServiceNow","Workday","Coupa","Kyriba"]),
           round(random.uniform(0.2,12),1), date_in("2026-06-01","2028-12-31"), random.choice(["Auto","Negotiate","Cancel"])] for i in range(40)]
    write_xlsx(p, {"Software": (["id","vendor","annual_m","renewal_date","renewal_action"], rows)}, title="Software Contracts", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Software Contracts", "Software license contracts.", sensitivity="Confidential")
    p = SRC / f / "bank_partner_contracts.xlsx"
    rows=[[f"BK-{i+1:02d}", bnk, random.choice(["Cash mgmt","Payments","FX","Debt"]), round(random.uniform(0.1,3.2),2),
           date_in("2026-06-01","2028-12-31")] for i,bnk in enumerate(BANKS)]
    write_xlsx(p, {"BankContracts": (["id","bank","service","annual_fee_m","renewal"], rows)}, title="Bank Partner Contracts", meta=M(f,"Elena Vasquez","Kyriba TMS"))
    reg(p, f, "Elena Vasquez", "Kyriba TMS", "Bank Partner Contracts", "Bank partner agreements.", sensitivity="Confidential")
    p = SRC / f / "rate_cards.xlsx"
    roles=["Partner","Director","Manager","Sr Consultant","Consultant","Analyst","Architect","Developer","Tester","PM"]
    rows=[]
    for prt in SI_PARTNERS:
        for r in roles:
            for loc in ["Onshore","Nearshore","Offshore"]:
                base={"Onshore":1.0,"Nearshore":0.62,"Offshore":0.38}[loc]
                rate={"Partner":420,"Director":340,"Architect":300,"Manager":260,"Sr Consultant":210,
                      "Consultant":170,"PM":230,"Developer":150,"Tester":120,"Analyst":110}[r]
                rows.append([prt, r, loc, round(rate*base)])
    rows=rows[:140]
    write_xlsx(p, {"RateCards": (["partner","role","location","hourly_usd"], rows)}, title="Rate Cards", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Rate Cards", "140-row SI rate card.", sensitivity="Confidential")
    p = SRC / f / "vendor_scorecards.xlsx"
    rows=[[prt, random.randint(60,98), random.randint(55,97), random.randint(50,99), random.choice(["A","B","C"])] for prt in SI_PARTNERS]
    write_xlsx(p, {"Scorecards": (["vendor","delivery_score","quality_score","value_score","tier"], rows)}, title="Vendor Scorecards", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Vendor Scorecards", "Vendor performance scorecards.")
    p = SRC / f / "sourcing_pipeline.xlsx"
    rows=[[f"SRC-{i+1:02d}", random.choice(["AMS RFP","Kyriba SI","Reporting tool","Cloud","Cyber"]),
           random.choice(["Intake","RFI","RFP","BAFO","Award"]), round(random.uniform(1,40),1)] for i in range(18)]
    write_xlsx(p, {"Pipeline": (["id","sourcing_event","stage","value_m"], rows)}, title="Sourcing Pipeline", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Sourcing Pipeline", "Active sourcing events.")
    p = SRC / f / "bafo_model.xlsx"
    rows=[[prt, round(random.uniform(8,40),1), round(random.uniform(7,36),1), f"{random.uniform(5,22):.1f}%"] for prt in SI_PARTNERS[:5]]
    write_xlsx(p, {"BAFO": (["bidder","initial_bid_m","bafo_bid_m","reduction"], rows)}, title="BAFO Model", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "BAFO Model", "Best-and-final-offer comparison.", sensitivity="Confidential")
    p = SRC / f / "contract_risks.xlsx"
    rows=[[f"CR-{i+1:02d}", random.choice(SI_PARTNERS), random.choice(["Auto-renew","Price escalation","Lock-in","SLA gap","Termination"]),
           random.choice(["High","Medium","Low"])] for i in range(20)]
    write_xlsx(p, {"ContractRisks": (["id","vendor","risk","severity"], rows)}, title="Contract Risks", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Contract Risks", "Contract risk register.")
    p = SRC / f / "renewal_calendar.xlsx"
    rows=[[date_in("2026-06-01","2027-12-31"), random.choice(SI_PARTNERS+["SAP","Oracle","Microsoft"]),
           round(random.uniform(0.2,12),1), random.choice(["Negotiate","Auto","Cancel"])] for _ in range(40)]
    rows.sort()
    write_xlsx(p, {"Renewals": (["renewal_date","vendor","annual_m","action"], rows)}, title="Renewal Calendar", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "Renewal Calendar", "Contract renewal calendar.")
    p = SRC / f / "rfp_requirements.docx"
    write_docx(p, "AMS RFP Requirements", [("Scope",[("bullets",["L2/L3 application support","Enhancements pool",
        "Kyriba run support","Reporting support","SLA & penalties","Continuous improvement"])]),
        ("Evaluation",[("table",["Criterion","Weight"],[["Capability","30%"],["Price","30%"],["Delivery model","20%"],["Innovation/AI","20%"]])])],
        meta=M(f,owner,sys))
    reg(p, f, owner, sys, "RFP Requirements", "AMS RFP requirements.")
    p = SRC / f / "ams_contract.pdf"
    write_pdf(p, "AMS Master Services Agreement (Synthetic)", "Application Managed Services agreement excerpt.",
        [("Key terms",[("table",["Term","Value"],
            [["Provider","Incumbent SI"],["Term","3 years + 1+1"],["Annual value","$18.4M"],
             ["SLA credits","Up to 12%"],["Benchmarking","Every 18 months"],["AI clause","Gain-share on automation"]])]),
         ("Risks",[("bullets",["Auto-renewal clause","Above-market rate card","Weak automation incentives"])])],
        meta=M(f,owner,sys))
    reg(p, f, owner, sys, "AMS Contract", "AMS master services agreement.", sensitivity="Confidential")

# ---- 11 risk, controls, audit, responsible AI -----------------------------
def build_risk():
    f = "11_risk_controls_responsible_ai"; owner="David Chen"
    p = SRC / f / "risk_control_register.xlsx"
    rows=[]
    for i in range(120):
        rows.append([f"RC-{i+1:04d}", random.choice(["SOX","Operational","Cyber","Treasury","Vendor","AI"]),
            random.choice(["Payment auth","Access review","SoD","Change mgmt","Backup","Vendor due diligence",
                "Model validation","Data privacy","Reconciliation","Monitoring"])+f" {i}",
            random.choice(["Preventive","Detective"]), random.choice(["Effective","Partial","Gap","Not tested"]),
            random.choice(["High","Medium","Low"]), random.choice(EXECS)[0]])
    write_xlsx(p, {"RiskControls": (["id","domain","control","type","effectiveness","risk_rating","owner"], rows)},
               title="Risk & Control Register", meta=M(f,owner,"GRC"))
    reg(p, f, owner, "GRC", "Risk & Control Register", "120-row master risk/control register.", sensitivity="Confidential")
    p = SRC / f / "sox_controls.xlsx"
    rows=[[f"SOX-{i+1:03d}", random.choice(["O2C","P2P","R2R","Treasury","ITGC"]), random.choice(["Key","Non-key"]),
           random.choice(["Effective","Deficiency","Material weakness candidate"])] for i in range(60)]
    write_xlsx(p, {"SOX": (["id","cycle","key_control","status"], rows)}, title="SOX Controls", meta=M(f,owner,"GRC"))
    reg(p, f, owner, "GRC", "SOX Controls", "SOX control inventory.", sensitivity="Confidential")
    p = SRC / f / "payment_fraud_controls.xlsx"
    rows=[[n, random.choice(["Preventive","Detective"]), random.choice(["Effective","Gap"])]
          for n in ["Dual authorization","Beneficiary validation","Sanctions screening","Payment limits","Anomaly detection","Callback verification"]]
    write_xlsx(p, {"FraudControls": (["control","type","status"], rows)}, title="Payment/Fraud Controls", meta=M(f,"Elena Vasquez","Kyriba TMS"))
    reg(p, f, "Elena Vasquez", "Kyriba TMS", "Payment/Fraud Controls", "Payment and fraud controls.", sensitivity="Confidential")
    p = SRC / f / "segregation_of_duties.xlsx"
    rows=[[f"SoD-{i+1:02d}", random.choice(["Create vendor + pay","Post + approve JE","Admin + user","Trade + settle"]),
           random.choice(["Conflict","Mitigated","Clean"]), random.choice(EXECS)[0]] for i in range(24)]
    write_xlsx(p, {"SoD": (["id","conflict_rule","status","owner"], rows)}, title="Segregation of Duties", meta=M(f,owner,"SailPoint"))
    reg(p, f, owner, "SailPoint", "Segregation of Duties", "SoD conflict matrix.", sensitivity="Confidential")
    p = SRC / f / "audit_findings.xlsx"
    rows=[[f"AF-{i+1:03d}", random.choice(["Internal","External","SOX"]), random.choice(["High","Medium","Low"]),
           random.choice(["Open","Remediating","Closed"]), date_in("2025-01-01","2026-06-01")] for i in range(45)]
    write_xlsx(p, {"AuditFindings": (["id","source","severity","status","raised"], rows)}, title="Audit Findings", meta=M(f,owner,"GRC"))
    reg(p, f, owner, "GRC", "Audit Findings", "Audit finding register.", sensitivity="Confidential")
    p = SRC / f / "cyber_risks.xlsx"
    rows=[[f"CY-{i+1:02d}", random.choice(["Phishing","Ransomware","Identity","Cloud misconfig","Third-party","Data exfil"]),
           random.choice(["High","Medium","Low"]), random.choice(["Open","Mitigating","Accepted"])] for i in range(20)]
    write_xlsx(p, {"CyberRisks": (["id","risk","severity","status"], rows)}, title="Cyber Risks", meta=M(f,owner,"Sentinel"))
    reg(p, f, owner, "Sentinel", "Cyber Risks", "Cyber risk register.", sensitivity="Confidential")
    p = SRC / f / "operational_risks.xlsx"
    rows=[[f"OR-{i+1:02d}", random.choice(["Process failure","Key person","Supplier","Compliance","Continuity"]),
           random.choice(["High","Medium","Low"])] for i in range(20)]
    write_xlsx(p, {"OpRisks": (["id","risk","severity"], rows)}, title="Operational Risks", meta=M(f,owner,"GRC"))
    reg(p, f, owner, "GRC", "Operational Risks", "Operational risk register.")
    p = SRC / f / "responsible_ai_controls.xlsx"
    rows=[[f"RAI-{i+1:02d}", n, random.choice(["Implemented","Partial","Planned"]), random.choice(EXECS)[0]]
          for i,n in enumerate(["Use-case intake","Bias testing","Human-in-loop","Explainability","Data lineage",
            "Model monitoring","Privacy review","Security review","Audit trail","Kill switch"])]
    write_xlsx(p, {"ResponsibleAI": (["id","control","status","owner"], rows)}, title="Responsible AI Controls", meta=M(f,"Aisha Bello","AI Governance"))
    reg(p, f, "Aisha Bello", "AI Governance", "Responsible AI Controls", "Responsible AI control set.")
    p = SRC / f / "model_use_case_review_checklist.xlsx"
    rows=[[q, random.choice(["Yes","No","Partial"])] for q in
          ["Business value defined?","Data sources approved?","PII assessed?","Bias tested?","Human oversight?",
           "Explainability documented?","Monitoring in place?","Security reviewed?","Audit trail enabled?","Rollback plan?"]]
    write_xlsx(p, {"Checklist": (["review_question","answer"], rows)}, title="Model/Use-Case Review Checklist", meta=M(f,"Aisha Bello","AI Governance"))
    reg(p, f, "Aisha Bello", "AI Governance", "Model Review Checklist", "AI model/use-case review checklist.")
    p = SRC / f / "evidence_gap_register.csv"
    rows=[[f"GAP-{i+1:03d}", random.choice(["Treasury","Finance","IT","Vendor","AI"]),
           random.choice(["No source doc","Stale evidence","Unverified claim","Owner unknown"]),
           random.choice(["High","Medium","Low"])] for i in range(40)]
    write_csv(p, ["gap_id","domain","gap_type","severity"], rows, meta=M(f,owner,"GRC"))
    reg(p, f, owner, "GRC", "Evidence Gap Register", "Evidence gaps for hardening.")
    p = SRC / f / "ai_governance.pdf"
    write_pdf(p, "AI Governance Framework", "Governed AI operating model and Responsible AI.",
        [("Framework",[("bullets",["Use-case intake & approval","Tiered risk classification","Human oversight by tier",
            "Model monitoring & drift","Evidence-cited outputs only"])]),
         ("Controls",[("table",["Tier","Oversight","Approval"],
            [["Low","Spot-check","CDO delegate"],["Medium","Human review","AI Board"],["High","Human-in-loop","AI Board + CRO"]])])],
        meta=M(f,"Aisha Bello","AI Governance"))
    reg(p, f, "Aisha Bello", "AI Governance", "AI Governance Framework", "AI governance and Responsible AI framework.")

print("built: vendors, risk")

# ---- 12 AI use cases & moves ----------------------------------------------
def build_ai():
    f = "12_ai_use_cases_moves"; owner="Marcus Reilly"; sys="AbarVa"
    usecases=["Cash forecasting copilot","Payment anomaly detection","Bank fee analyzer","Close acceleration assistant",
        "Reporting rationalization agent","Spend optimization agent","Contract risk extractor","Vendor scorecard agent",
        "Incident summarizer","Knowledge deflection bot","SoD conflict detector","Audit evidence assembler",
        "Forecast variance explainer","Liquidity scenario modeler","FX hedge advisor","Working capital optimizer"]
    p = SRC / f / "ai_opportunity_portfolio.xlsx"
    rows=[]
    for i in range(60):
        uc=random.choice(usecases)
        rows.append([f"AI-{i+1:03d}", uc+f" v{i%4}", random.choice(["Finance","Treasury","IT","Procurement","Risk","Data"]),
            round(random.uniform(0.4,9.5),1), random.choice(["High","Medium","Low"]),
            random.choice(["Quick win","Strategic","Foundational"]), random.choice(["Backlog","Pilot","Scaling","Live"]),
            random.choice(["Low","Medium","High"])])
    write_xlsx(p, {"AIPortfolio": (["id","use_case","domain","value_m","value_confidence","type","status","ai_risk_tier"], rows)},
               title="AI Opportunity Portfolio", meta=M(f,owner,sys))
    reg(p, f, owner, sys, "AI Opportunity Portfolio", "60-row enterprise AI opportunity portfolio.")
    for fn,title,dom in [("finance_ai_use_cases.xlsx","Finance AI Use Cases","Finance"),
                         ("it_modernization_use_cases.xlsx","IT Modernization Use Cases","IT"),
                         ("procurement_vendor_optimization_use_cases.xlsx","Procurement/Vendor Optimization Use Cases","Procurement"),
                         ("cost_optimization_use_case.xlsx","Cost Optimization Use Cases","Finance"),
                         ("value_realization_model.xlsx","Value Realization Model","Finance")]:
        p = SRC / f / fn
        rows=[[f"UC-{i+1:02d}", random.choice(usecases), round(random.uniform(0.3,8),1),
               random.choice(["6m","12m","18m"]), random.choice(["High","Medium"])] for i in range(14)]
        write_xlsx(p, {dom[:18]: (["id","use_case","benefit_m","horizon","confidence"], rows)}, title=title, meta=M(f,owner,sys))
        reg(p, f, owner, sys, title, f"{title}.")
    p = SRC / f / "treasury_kyriba_success_use_case.pdf"
    write_pdf(p, "Treasury / Kyriba Success Use Case", "Reference success pattern for treasury AI.",
        [("Outcome",[("bullets",["Cash visibility within 48h of go-live","$160M idle cash released","32% fewer payment exceptions"])]),
         ("Enablers",[("bullets",["Bank connectivity standardization","AI cash forecasting","Anomaly detection on payments"])])],
        meta=M(f,"Elena Vasquez","Kyriba TMS"))
    reg(p, f, "Elena Vasquez", "Kyriba TMS", "Treasury/Kyriba Success Use Case", "Reference success pattern.")
    p = SRC / f / "reporting_rationalization_use_case.docx"
    write_docx(p, "Reporting Rationalization Use Case", [("Approach",[("bullets",[
        "AI-assisted usage analysis","Duplicate metric detection","Auto-generated semantic mappings","Self-serve certified metrics"])]),
        ("Value",["55% manual reporting effort reduction; 320 → 140 managed reports."])], meta=M(f,"Aisha Bello","Power BI"))
    reg(p, f, "Aisha Bello", "Power BI", "Reporting Rationalization Use Case", "Reporting rationalization AI use case.")
    p = SRC / f / "human_agent_operating_model.docx"
    write_docx(p, "Human + Agent Operating Model", [("Principles",[("bullets",[
        "Agents draft, humans decide","Evidence-cited outputs only","Tiered autonomy by risk",
        "Continuous monitoring & audit trail"])]),
        ("Roles",[("table",["Role","Responsibility"],[["Domain expert","Approves outputs"],
            ["AI steward","Monitors models"],["Risk owner","Sets guardrails"]])])], meta=M(f,owner,"AbarVa"))
    reg(p, f, owner, "AbarVa", "Human + Agent Operating Model", "Human/agent operating model.")
    p = SRC / f / "ai_moves_portfolio_summary.md"
    write_md(p, "AI Moves Portfolio Summary",
        "## Top moves\n\n" + "\n".join(f"{i+1}. **{u}**" for i,u in enumerate(usecases[:8])) +
        "\n\n## Flagship\nLakeshore Enterprise Finance & Treasury Modernization (Kyriba rollout, corporate controls, "
        "reporting rationalization, vendor optimization, value realization).",
        meta=M(f,owner,"AbarVa"))
    reg(p, f, owner, "AbarVa", "AI Moves Portfolio Summary", "Summary of AI Moves portfolio.")

# ============================================================================
# SVG ARCHITECTURE TOOLKIT — real, technical diagrams
# ============================================================================
PAL = {
    "ink": "#0c1a3a", "navy": "#1F2A44", "paper": "#f5f1eb", "card": "#ffffff",
    "line": "#26324f", "muted": "#6b7280", "edge": "#39507e",
    "z_channel": "#eef2fb", "z_edge": "#fdeef0", "z_app": "#eef7f1",
    "z_int": "#fff7e8", "z_data": "#eef4fb", "z_cloud": "#f1eefb",
    "z_sec": "#fdeef0", "z_obs": "#eefbf7",
    "ac_blue": "#0066CC", "ac_teal": "#0f766e", "ac_amber": "#b45309",
    "ac_red": "#b91c1c", "ac_purple": "#6d28d9", "ac_green": "#15803d",
}

def esc(s):
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def _svg_header(width, height, title, meta, subtitle=""):
    title = esc(title); subtitle = esc(subtitle)
    return (
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" '
        f'viewBox="0 0 {width} {height}" font-family="Inter, Helvetica, Arial, sans-serif">\n'
        f'<defs>\n'
        f'<marker id="arw" markerWidth="9" markerHeight="9" refX="7" refY="3" orient="auto" markerUnits="strokeWidth">'
        f'<path d="M0,0 L7,3 L0,6 z" fill="{PAL["edge"]}"/></marker>\n'
        f'<marker id="arwR" markerWidth="9" markerHeight="9" refX="7" refY="3" orient="auto" markerUnits="strokeWidth">'
        f'<path d="M0,0 L7,3 L0,6 z" fill="{PAL["ac_red"]}"/></marker>\n'
        f'<marker id="dot" markerWidth="6" markerHeight="6" refX="3" refY="3"><circle cx="3" cy="3" r="2.4" fill="{PAL["edge"]}"/></marker>\n'
        f'</defs>\n'
        f'<rect width="{width}" height="{height}" fill="{PAL["paper"]}"/>\n'
        f'<text x="36" y="34" font-size="12" fill="{PAL["muted"]}" font-family="monospace">{WATERMARK}</text>\n'
        f'<text x="36" y="64" font-size="23" font-weight="700" fill="{PAL["ink"]}" font-family="Georgia, serif">{title}</text>\n'
        + (f'<text x="36" y="86" font-size="12.5" fill="{PAL["muted"]}">{subtitle}</text>\n' if subtitle else '')
        + f'<text x="{width/2}" y="{height/2}" font-size="150" fill="{PAL["ink"]}" opacity="0.045" '
          f'text-anchor="middle" transform="rotate(-22 {width/2} {height/2})" font-weight="700">SYNTHETIC</text>\n'
    )

def _svg_footer(width, height, meta):
    return (f'<text x="36" y="{height-14}" font-size="10.5" fill="{PAL["muted"]}" font-family="monospace">'
            f'{PACK_ID} · {meta["domain"]} · owner={meta["owner"]} · system={meta["system"]} · {meta["date"]} · {WATERMARK_SHORT}</text>\n</svg>\n')

def _zone(x, y, w, h, label, fill):
    label = esc(label)
    return (f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="12" fill="{fill}" stroke="{PAL["line"]}" '
            f'stroke-opacity="0.18"/>\n'
            f'<text x="{x+4}" y="{y-5}" font-size="11" font-weight="700" letter-spacing="1.5" '
            f'fill="{PAL["navy"]}" font-family="monospace">{label}</text>\n')

def _node(x, y, w, h, title, subs=None, accent=None, mono_tag=None):
    accent = accent or PAL["navy"]
    s = (f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="9" fill="{PAL["card"]}" stroke="{PAL["line"]}" stroke-width="1.1"/>\n'
         f'<rect x="{x}" y="{y}" width="5" height="{h}" rx="2" fill="{accent}"/>\n'
         f'<text x="{x+14}" y="{y+21}" font-size="12.5" font-weight="600" fill="{PAL["ink"]}">{esc(title)}</text>\n')
    if mono_tag:
        s += (f'<text x="{x+w-10}" y="{y+18}" font-size="8.5" fill="{accent}" text-anchor="end" '
              f'font-family="monospace" font-weight="700">{esc(mono_tag)}</text>\n')
    for i, ln in enumerate(subs or []):
        s += f'<text x="{x+14}" y="{y+40+i*15}" font-size="10.5" fill="{PAL["muted"]}">{esc(ln)}</text>\n'
    return s

def _edge(x1, y1, x2, y2, label="", dashed=False, color=None, red=False, curve=0):
    color = PAL["ac_red"] if red else (color or PAL["edge"])
    mk = "url(#arwR)" if red else "url(#arw)"
    dash = 'stroke-dasharray="5 4" ' if dashed else ''
    if curve:
        mx = (x1 + x2) / 2
        path = f'M{x1},{y1} C{mx},{y1+curve} {mx},{y2-curve} {x2},{y2}'
    else:
        path = f'M{x1},{y1} L{x2},{y2}'
    s = f'<path d="{path}" fill="none" stroke="{color}" stroke-width="1.6" {dash}marker-end="{mk}"/>\n'
    if label:
        lx, ly = (x1 + x2) / 2, (y1 + y2) / 2
        w = len(label) * 5.4 + 10
        s += (f'<rect x="{lx-w/2}" y="{ly-9}" width="{w}" height="15" rx="3" fill="{PAL["paper"]}" opacity="0.92"/>'
              f'<text x="{lx}" y="{ly+2}" font-size="9" fill="{color}" text-anchor="middle" font-family="monospace">{esc(label)}</text>\n')
    return s

def _legend(x, y, items):
    s = f'<rect x="{x}" y="{y}" width="250" height="{18+len(items)*16}" rx="8" fill="{PAL["card"]}" stroke="{PAL["line"]}" stroke-opacity="0.3"/>\n'
    s += f'<text x="{x+12}" y="{y+16}" font-size="10" font-weight="700" font-family="monospace" fill="{PAL["navy"]}">LEGEND</text>\n'
    for i, (c, t) in enumerate(items):
        yy = y + 30 + i * 16
        s += (f'<rect x="{x+12}" y="{yy-8}" width="14" height="9" rx="2" fill="{c}"/>'
              f'<text x="{x+32}" y="{yy}" font-size="9.5" fill="{PAL["muted"]}">{esc(t)}</text>\n')
    return s

def write_arch_svg(path, title, subtitle, width, height, inner, meta):
    svg = _svg_header(width, height, title, meta, subtitle) + inner + _svg_footer(width, height, meta)
    path.write_text(svg, encoding="utf-8")

# ---- individual diagrams --------------------------------------------------
def diagram_enterprise_current_state():
    W, H = 1320, 900
    inner = ""
    inner += _zone(36, 110, 1248, 78, "EXPERIENCE / CHANNELS", PAL["z_channel"])
    chs = [("Employee Web", "SSO · HTTPS"), ("Mobile (iOS/Android)", "MDM · OAuth2"),
           ("Supplier Portal", "Coupa SAN"), ("Customer EDI", "AS2 · X12"), ("Bank Portals", "2FA")]
    for i, (t, s) in enumerate(chs):
        inner += _node(52 + i * 246, 128, 226, 48, t, [s], PAL["ac_blue"])
    inner += _zone(36, 204, 1248, 70, "EDGE / IDENTITY / SECURITY", PAL["z_edge"])
    edge = [("Azure Front Door + WAF", "TLS1.2 · DDoS"), ("Entra ID (Azure AD)", "SAML/OIDC · SCIM"),
            ("Zscaler SSE", "ZTNA"), ("CyberArk PAM", "JIT secrets"), ("API Mgmt Gateway", "OAuth2 · rate-limit")]
    for i, (t, s) in enumerate(edge):
        inner += _node(52 + i * 246, 220, 226, 46, t, [s], PAL["ac_red"])
    inner += _zone(36, 290, 1248, 196, "CORE APPLICATION ESTATE", PAL["z_app"])
    apps = [
        ("SAP ECC 6.0 (NA/EU)", ["GL/AP/AR/MM/SD", "IDoc · BAPI · RFC"], PAL["ac_teal"], "ERP"),
        ("SAP S/4HANA (APAC)", ["Finance · OData", "migration wave 2"], PAL["ac_teal"], "ERP"),
        ("Workday HCM", ["HR · payroll", "REST · EIB"], PAL["ac_purple"], "HRIS"),
        ("Coupa", ["P2P · sourcing", "cXML · REST"], PAL["ac_amber"], "P2P"),
        ("Kyriba TMS", ["cash · payments", "ISO20022 · API"], PAL["ac_blue"], "TMS"),
        ("ServiceNow", ["ITSM · CMDB", "REST · MID"], PAL["ac_green"], "ITSM"),
        ("Salesforce", ["O2C · CPQ", "Bulk API"], PAL["ac_blue"], "CRM"),
        ("Hyperion/OneStream", ["consolidation · FP&A", "Smart View"], PAL["ac_teal"], "EPM"),
    ]
    for i, (t, s, a, tag) in enumerate(apps):
        col = i % 4; row = i // 4
        inner += _node(52 + col * 312, 312 + row * 86, 292, 74, t, s, a, tag)
    inner += _zone(36, 502, 1248, 96, "INTEGRATION & EVENTING", PAL["z_int"])
    integ = [("SAP PI/PO (ESB)", ["IDoc/SOAP", "A2A · B2B"], "ESB"),
             ("Azure API Mgmt", ["REST facades", "policy/quotas"], "APIM"),
             ("Azure Service Bus", ["q-context-*", "pub/sub"], "EVENT"),
             ("MuleSoft iPaaS", ["SaaS connectors", "batch+rt"], "IPAAS"),
             ("SFTP / MFT", ["bank files", "PGP"], "MFT")]
    for i, (t, s, tag) in enumerate(integ):
        inner += _node(52 + i * 246, 520, 226, 70, t, s, PAL["ac_amber"], tag)
    inner += _zone(36, 614, 858, 150, "DATA & ANALYTICS PLATFORM (Azure + Snowflake)", PAL["z_data"])
    inner += _node(52, 636, 196, 120, "Snowflake", ["RAW → CURATED →", "PRESENT (medallion)", "Streams · Tasks", "RBAC · masking"], PAL["ac_blue"], "DW")
    inner += _node(258, 636, 196, 120, "dbt + Semantic Layer", ["certified metrics", "tests · lineage", "exposures"], PAL["ac_teal"], "ELT")
    inner += _node(464, 636, 196, 120, "Azure Data Lake", ["landing/bronze", "Parquet · ADLS g2"], PAL["ac_purple"], "LAKE")
    inner += _node(670, 636, 208, 120, "Power BI / AI Search", ["governed datasets", "tenant-context-v1", "vector + BM25"], PAL["ac_green"], "BI/RAG")
    inner += _zone(912, 614, 372, 150, "PLATFORM / OBSERVABILITY", PAL["z_obs"])
    inner += _node(928, 636, 168, 120, "Azure Postgres", ["enterprise_context_*", "control plane", "PITR · HA"], PAL["ac_blue"], "OLTP")
    inner += _node(1104, 636, 164, 120, "Azure Monitor", ["App Insights", "Log Analytics", "Defender"], PAL["ac_amber"], "OBS")
    # data-flow edges with protocols
    inner += _edge(165, 176, 165, 220, red=True)  # channel->edge
    inner += _edge(165, 266, 198, 312, "OIDC", color=PAL["ac_red"])
    inner += _edge(198, 386, 198, 520, "IDoc", dashed=True)            # SAP->PI/PO
    inner += _edge(510, 386, 510, 520, "cXML")                         # Coupa->APIM
    inner += _edge(822, 386, 700, 520, "REST")                         # SFNow->ServiceBus
    inner += _edge(1134, 386, 600, 520, "ISO20022", color=PAL["ac_blue"])  # Kyriba->MFT
    inner += _edge(165, 590, 150, 636, "batch")                        # PI/PO->Snowflake
    inner += _edge(411, 590, 356, 636, "CDC")                          # APIM->dbt
    inner += _edge(150, 696, 670, 696, "metrics", dashed=True)         # snowflake->BI
    inner += _edge(880, 690, 928, 690, "JDBC")                         # data->postgres
    inner += _legend(1040, 110, [(PAL["ac_red"], "identity / security flow"),
                                 (PAL["edge"], "synchronous API/data flow"),
                                 (PAL["ac_blue"], "treasury / payment (ISO20022)")])
    return W, H, inner

def diagram_integration():
    W, H = 1300, 820
    inner = _node  # alias unused
    s = ""
    s += _zone(36, 110, 1228, 120, "SYSTEMS OF RECORD (producers)", PAL["z_app"])
    src = [("SAP ECC/S4", "IDoc/BAPI"), ("Workday", "EIB/REST"), ("Coupa", "cXML"),
           ("Kyriba", "API/ISO20022"), ("ServiceNow", "REST"), ("Salesforce", "Bulk API")]
    for i, (t, p) in enumerate(src):
        s += _node(52 + i * 202, 134, 186, 80, t, [p, "real-time + batch"], PAL["ac_teal"])
    s += _zone(36, 256, 1228, 150, "INTEGRATION BACKBONE", PAL["z_int"])
    s += _node(60, 282, 270, 110, "SAP PI/PO Enterprise Service Bus", ["A2A orchestration, mapping (XSLT)", "guaranteed delivery, monitoring", "240+ active interfaces"], PAL["ac_amber"], "ESB")
    s += _node(350, 282, 270, 110, "Azure API Management", ["REST facade + OAuth2 policies", "rate limiting, versioning", "developer portal"], PAL["ac_blue"], "APIM")
    s += _node(640, 282, 270, 110, "Azure Service Bus / Event Grid", ["q-context-ingestion-events", "pub/sub, dead-letter, sessions", "competing consumers"], PAL["ac_purple"], "EVENT")
    s += _node(930, 282, 310, 110, "Managed File Transfer (SFTP/MFT)", ["bank host-to-host, PGP encryption", "pain.001 out / pain.002,camt.05x in", "scheduled + event-driven"], PAL["ac_green"], "MFT")
    s += _zone(36, 432, 1228, 120, "CONSUMERS (subscribers)", PAL["z_data"])
    con = [("Snowflake (ELT)", "Snowpipe/COPY"), ("Power BI", "DirectQuery"), ("AI Search", "indexers"),
           ("Azure Postgres", "JDBC"), ("Partner Banks", "ISO20022"), ("Context Loader", "Service Bus")]
    for i, (t, p) in enumerate(con):
        s += _node(52 + i * 202, 456, 186, 80, t, [p, "idempotent sink"], PAL["ac_blue"])
    # flows
    for i in range(6):
        s += _edge(145 + i * 202, 214, 195 + (i % 4) * 290, 282, "", dashed=(i % 2 == 0))
    s += _edge(195, 392, 145, 456, "ELT")
    s += _edge(485, 392, 700, 456, "events")
    s += _edge(1085, 392, 1055, 456, "ISO20022", color=PAL["ac_blue"])
    s += _zone(36, 576, 1228, 168, "INTERFACE PATTERN INVENTORY (illustrative)", PAL["z_obs"])
    rows = [["IF-ERP-TMS-01", "SAP→Kyriba", "Payment proposals", "ISO20022 pain.001", "event", "hourly"],
            ["IF-TMS-BANK-02", "Kyriba→Bank", "Outbound payments", "pain.001.001.09", "MFT", "intraday"],
            ["IF-BANK-TMS-03", "Bank→Kyriba", "Statements", "camt.053/052", "MFT", "daily/intraday"],
            ["IF-ERP-DW-04", "SAP→Snowflake", "GL/AP/AR extract", "IDoc→Parquet", "batch", "nightly"],
            ["IF-SNOW-CTX-05", "Snowflake→Postgres", "Context chunks", "JDBC upsert", "batch", "on-load"],
            ["IF-CTX-SEARCH-06", "Postgres→AI Search", "Index refresh", "REST index API", "event", "on-commit"]]
    hdr = ["interface_id", "route", "payload", "format/standard", "pattern", "frequency"]
    colw = [150, 170, 200, 210, 130, 180]
    x0 = 60
    cx = x0
    for j, h in enumerate(hdr):
        s += f'<text x="{cx}" y="{612}" font-size="10.5" font-weight="700" font-family="monospace" fill="{PAL["navy"]}">{h}</text>'
        cx += colw[j]
    for r_i, row in enumerate(rows):
        cx = x0; yy = 632 + r_i * 18
        for j, val in enumerate(row):
            s += f'<text x="{cx}" y="{yy}" font-size="10" fill="{PAL["ink"]}">{val}</text>'
            cx += colw[j]
    return W, H, s

def diagram_kyriba():
    W, H = 1300, 840
    s = ""
    s += _zone(36, 110, 600, 150, "ERP / SOURCE", PAL["z_app"])
    s += _node(56, 134, 270, 110, "SAP ECC / S/4HANA", ["AP proposals, GL postings", "vendor master, bank master", "IDoc PEXR2002 / OData"], PAL["ac_teal"], "ERP")
    s += _node(346, 134, 270, 110, "Coupa / Concur", ["invoice approvals, expenses", "supplier bank details"], PAL["ac_amber"], "P2P")
    s += _zone(660, 110, 604, 150, "TREASURY MANAGEMENT (Kyriba SaaS)", PAL["z_data"])
    s += _node(680, 134, 270, 110, "Kyriba Payments Factory", ["payment workflow + limits", "dual auth, sanctions screen", "format transformation"], PAL["ac_blue"], "PAY")
    s += _node(970, 134, 274, 110, "Kyriba Cash & Liquidity", ["cash positioning, pooling", "13-week forecast (AI)", "bank fee analysis"], PAL["ac_blue"], "CASH")
    s += _zone(36, 300, 1228, 130, "CONNECTIVITY LAYER", PAL["z_int"])
    s += _node(60, 322, 280, 92, "Kyriba Bank Connectivity", ["SWIFT Alliance Lite2 / SCORE", "host-to-host SFTP, EBICS, API", "format library ISO20022"], PAL["ac_purple"], "CONN")
    s += _node(360, 322, 280, 92, "Sanctions / AML", ["real-time screening", "OFAC/EU lists", "hold & review queue"], PAL["ac_red"], "AML")
    s += _node(660, 322, 280, 92, "ERP Reconciliation", ["auto-recon statements↔GL", "exception workflow", "fee validation"], PAL["ac_green"], "RECON")
    s += _node(960, 322, 284, 92, "Treasury Controls", ["segregation of duties", "audit trail (immutable)", "limit & mandate mgmt"], PAL["ac_amber"], "CTRL")
    s += _zone(36, 470, 1228, 150, "BANKING PARTNERS (10 banks · 4 regions)", PAL["z_channel"])
    banks = ["JPMorgan", "Citi", "HSBC", "BNP Paribas", "Std Chartered", "BofA", "Deutsche", "DBS", "Santander", "Wells Fargo"]
    for i, b in enumerate(banks):
        col = i % 5; row = i // 5
        s += _node(56 + col * 240, 494 + row * 58, 224, 48, b, ["camt.053 ⇄ pain.001/002"], PAL["ac_blue"])
    # flow chain
    s += _edge(326, 188, 680, 188, "pain.001", color=PAL["ac_blue"])
    s += _edge(815, 244, 200, 322, "proposals", dashed=True)
    s += _edge(200, 414, 200, 494, "host-to-host", color=PAL["ac_blue"])
    s += _edge(500, 414, 500, 494, "screened", red=True)
    s += _edge(500, 552, 800, 414, "camt.05x", curve=40)
    s += _edge(800, 414, 1100, 244, "recon → cash", dashed=True, curve=-30)
    s += _zone(36, 632, 1228, 120, "PAYMENT MESSAGE FLOW (end-to-end)", PAL["z_obs"])
    steps = ["1. ERP raises AP proposal", "2. Kyriba enriches + validates limits", "3. dual authorization",
             "4. sanctions screening", "5. ISO20022 pain.001 to bank (SFTP/SWIFT)", "6. bank ack pain.002",
             "7. statement camt.053 inbound", "8. auto-reconciliation to GL", "9. cash position + forecast refresh"]
    for i, st in enumerate(steps):
        col = i % 3; row = i // 3
        x = 60 + col * 400; y = 668 + row * 26
        s += f'<text x="{x}" y="{y}" font-size="11" fill="{PAL["ink"]}">{st}</text>'
    return W, H, s

def diagram_data_platform():
    W, H = 1300, 820
    s = ""
    s += _zone(36, 110, 1228, 110, "SOURCES", PAL["z_app"])
    for i, (t, p) in enumerate([("SAP", "IDoc/CDC"), ("Workday", "EIB"), ("Coupa", "cXML"), ("Kyriba", "API"),
                                 ("ServiceNow", "REST"), ("Salesforce", "Bulk")]):
        s += _node(52 + i * 202, 132, 186, 72, t, [p], PAL["ac_teal"])
    s += _zone(36, 244, 1228, 250, "SNOWFLAKE MEDALLION + GOVERNANCE", PAL["z_data"])
    s += _node(60, 270, 360, 96, "BRONZE / RAW (ADLS gen2 + ext stages)", ["Snowpipe auto-ingest, Parquet", "immutable landing, schema-on-read", "PII tagging at ingest"], PAL["ac_purple"], "BRONZE")
    s += _node(60, 376, 360, 96, "SILVER / CURATED (dbt models)", ["conformed dims, SCD2 history", "data quality tests (dbt/Great Exp.)", "deduped, standardized"], PAL["ac_blue"], "SILVER")
    s += _node(440, 270, 360, 96, "GOLD / PRESENT (marts)", ["finance, treasury, vendor, ITSM", "aggregates + exposures", "row access policies"], PAL["ac_green"], "GOLD")
    s += _node(440, 376, 360, 96, "SEMANTIC LAYER (dbt metrics)", ["certified KPI definitions", "single source for BI + AI", "lineage + ownership"], PAL["ac_teal"], "SEMANTIC")
    s += _node(820, 270, 420, 96, "GOVERNANCE", ["RBAC + dynamic data masking", "object tagging, access history", "data contracts, SLAs"], PAL["ac_amber"], "GOV")
    s += _node(820, 376, 420, 96, "DATA QUALITY & OBSERVABILITY", ["freshness/volume/anomaly monitors", "incident → ServiceNow", "DQ scorecards"], PAL["ac_red"], "DQ")
    s += _zone(36, 516, 1228, 110, "CONSUMPTION", PAL["z_channel"])
    for i, (t, p, a) in enumerate([("Power BI (certified)", "governed datasets", PAL["ac_green"]),
                                    ("Azure AI Search", "tenant-context-v1 (vector+BM25)", PAL["ac_blue"]),
                                    ("Context Layer (Postgres)", "enterprise_context_chunks", PAL["ac_purple"]),
                                    ("Reverse ETL", "ops activation", PAL["ac_amber"])]):
        s += _node(52 + i * 306, 540, 288, 72, t, [p], a)
    # flows
    for i in range(6):
        s += _edge(145 + i * 202, 204, 240, 270, "", dashed=True)
    s += _edge(240, 366, 240, 376, "")
    s += _edge(420, 424, 440, 424, "test→")
    s += _edge(620, 472, 200, 540, "datasets", curve=30)
    s += _edge(620, 472, 500, 540, "embed→index", color=PAL["ac_blue"], curve=20)
    s += _edge(620, 472, 800, 540, "chunks", color=PAL["ac_purple"], curve=20)
    s += _legend(1010, 632, [(PAL["ac_purple"], "raw/landing"), (PAL["ac_blue"], "curated/vector"), (PAL["ac_green"], "gold/BI")])
    return W, H, s

def diagram_security():
    W, H = 1300, 800
    s = ""
    s += _zone(36, 110, 1228, 110, "IDENTITY-FIRST PERIMETER (Zero Trust)", PAL["z_edge"])
    for i, (t, p, a) in enumerate([("Entra ID", "SAML/OIDC, MFA, CA policies", PAL["ac_red"]),
                                    ("Zscaler SSE/ZTNA", "no implicit trust", PAL["ac_red"]),
                                    ("Front Door + WAF", "OWASP, DDoS", PAL["ac_amber"]),
                                    ("API Mgmt", "OAuth2, mTLS", PAL["ac_blue"]),
                                    ("CyberArk PAM", "JIT, vaulting", PAL["ac_purple"])]):
        s += _node(52 + i * 246, 134, 226, 72, t, [p], a)
    s += _zone(36, 244, 1228, 150, "PRIVATE DATA PLANE (VNet-isolated, private endpoints)", PAL["z_data"])
    for i, (t, p, a) in enumerate([("Azure Postgres", "private endpoint, AAD auth, TLS", PAL["ac_blue"]),
                                    ("Blob Storage", "deny-public, firewall+PE, CMK", PAL["ac_purple"]),
                                    ("Azure AI Search", "RBAC data-plane, PE", PAL["ac_green"]),
                                    ("Key Vault", "secrets, RBAC, soft-delete/purge", PAL["ac_amber"])]):
        s += _node(52 + i * 306, 268, 288, 100, t, [p, "managed identity access", "diagnostic logs → Sentinel"], a)
    s += _zone(36, 416, 1228, 130, "DETECT & RESPOND", PAL["z_obs"])
    for i, (t, p, a) in enumerate([("Microsoft Sentinel (SIEM)", "analytics rules, UEBA", PAL["ac_red"]),
                                    ("Defender for Cloud", "CSPM/CWPP, malware scan", PAL["ac_amber"]),
                                    ("Purview", "data classification, DLP", PAL["ac_blue"]),
                                    ("SOAR / Playbooks", "auto-containment", PAL["ac_green"])]):
        s += _node(52 + i * 306, 440, 288, 86, t, [p, "evidence → audit ledger"], a)
    s += _zone(36, 568, 1228, 150, "GOVERNANCE & CONTROLS MAPPING", PAL["z_app"])
    rows = [["SOX ITGC", "access review, change mgmt, SoD", "quarterly attestation"],
            ["Payment fraud", "dual auth, sanctions, anomaly", "Kyriba + Sentinel"],
            ["Data privacy", "masking, DLP, residency", "Purview"],
            ["Responsible AI", "use-case intake, human-in-loop", "AI Governance Board"]]
    for r_i, row in enumerate(rows):
        yy = 600 + r_i * 26
        s += f'<text x="60" y="{yy}" font-size="11" font-weight="600" fill="{PAL["ink"]}">{row[0]}</text>'
        s += f'<text x="240" y="{yy}" font-size="11" fill="{PAL["muted"]}">{row[1]}</text>'
        s += f'<text x="760" y="{yy}" font-size="11" fill="{PAL["navy"]}" font-family="monospace">{row[2]}</text>'
    s += _edge(165, 206, 200, 268, "MI", red=True)
    s += _edge(200, 368, 200, 440, "logs", dashed=True)
    return W, H, s

def diagram_ai_reference():
    W, H = 1300, 800
    s = ""
    s += _zone(36, 110, 1228, 96, "EXPERIENCE (governed surfaces)", PAL["z_channel"])
    for i, (t, p) in enumerate([("Sentinel Advisor (Ask)", "cited Q&A"), ("Nexus Sessions", "guided gates"),
                                 ("Strategic Moves", "board artifacts"), ("Control Tower", "value tracking")]):
        s += _node(52 + i * 306, 132, 288, 60, t, [p], PAL["ac_blue"])
    s += _zone(36, 220, 1228, 150, "ORCHESTRATION & GUARDRAILS", PAL["z_int"])
    s += _node(60, 244, 380, 110, "Agent Orchestrator", ["planner + tools, retries", "human-in-loop by risk tier", "cost + latency budgets"], PAL["ac_amber"], "ORCH")
    s += _node(450, 244, 380, 110, "Responsible AI Guardrails", ["PII redaction, grounding check", "evidence-required outputs", "policy: no uncited claims"], PAL["ac_red"], "RAI")
    s += _node(840, 244, 400, 110, "Evidence & Audit Ledger", ["every claim → source chunk", "immutable, replayable", "evidence_ledger (Postgres)"], PAL["ac_purple"], "AUDIT")
    s += _zone(36, 384, 1228, 160, "RETRIEVAL (RAG over enterprise context)", PAL["z_data"])
    s += _node(60, 408, 360, 116, "Context Broker", ["hybrid retrieve (vector+BM25)", "tenant-scoped (client_id)", "rerank + dedupe"], PAL["ac_teal"], "BROKER")
    s += _node(440, 408, 360, 116, "Azure AI Search", ["tenant-context-v1", "HNSW vector + semantic", "filter: tenant_key"], PAL["ac_blue"], "INDEX")
    s += _node(820, 408, 420, 116, "Azure Postgres context", ["enterprise_context_chunks", "embeddings (1536-d)", "provenance + metadata"], PAL["ac_green"], "STORE")
    s += _zone(36, 556, 1228, 110, "MODELS & INGESTION", PAL["z_obs"])
    s += _node(60, 580, 380, 72, "LLM / Embeddings", ["Claude (reasoning) + OpenAI emb", "text-embedding-3-small 1536-d"], PAL["ac_purple"], "MODEL")
    s += _node(450, 580, 380, 72, "Setup Admin Loader", ["ZIP→Blob→parse→chunk", "embed→index", "stage_and_process"], PAL["ac_amber"], "LOAD")
    s += _node(840, 580, 400, 72, "Document Intelligence", ["pdf/docx/xlsx parse", "layout + tables"], PAL["ac_blue"], "PARSE")
    s += _edge(200, 192, 250, 244, "query")
    s += _edge(250, 354, 240, 408, "retrieve")
    s += _edge(420, 466, 440, 466, "hybrid")
    s += _edge(800, 466, 820, 466, "vectors")
    s += _edge(640, 652, 240, 524, "embed", color=PAL["ac_purple"], curve=-30)
    s += _edge(640, 580, 1030, 524, "index", color=PAL["ac_blue"], curve=-30)
    s += _edge(900, 300, 1040, 192, "cite", dashed=True, color=PAL["ac_purple"])
    return W, H, s

def rich_diagrams():
    specs = [
        ("06_it_systems_architecture", "current_state_architecture.svg", diagram_enterprise_current_state,
         "Current-State Enterprise Architecture", "Channels → identity/security → core apps → integration → data → platform",
         "Liam O'Sullivan", "Architecture", "Layered enterprise reference architecture with protocol-labeled data flows."),
        ("06_it_systems_architecture", "integration_architecture_diagram.svg", diagram_integration,
         "Integration & Eventing Architecture", "Producers, backbone (ESB/APIM/Service Bus/MFT), consumers, interface inventory",
         "Liam O'Sullivan", "SAP PI/PO", "Integration backbone with interface pattern inventory."),
        ("05_treasury_kyriba", "kyriba_connectivity_architecture.svg", diagram_kyriba,
         "Kyriba Treasury Connectivity Architecture", "ERP → Kyriba payments/cash → bank connectivity → 10 banking partners",
         "Elena Vasquez", "Kyriba TMS", "End-to-end treasury & ISO20022 payment message architecture."),
        ("07_data_analytics_reporting", "data_platform_architecture.svg", diagram_data_platform,
         "Data & Analytics Platform Architecture", "Sources → Snowflake medallion + semantic layer + governance → consumption",
         "Aisha Bello", "Snowflake", "Medallion data platform with governance, DQ, and AI/BI consumption."),
        ("11_risk_controls_responsible_ai", "security_zero_trust_architecture.svg", diagram_security,
         "Security & Zero-Trust Reference Architecture", "Identity perimeter → private data plane → detect/respond → controls mapping",
         "David Chen", "Microsoft Sentinel", "Zero-trust security architecture with controls mapping."),
        ("12_ai_use_cases_moves", "ai_agent_reference_architecture.svg", diagram_ai_reference,
         "Enterprise AI & Agent Reference Architecture", "Governed surfaces → orchestration/guardrails → RAG retrieval → models/ingestion",
         "Marcus Reilly", "AbarVa", "Reference architecture for governed AI/agents over the enterprise context layer."),
    ]
    for folder, fn, builder, title, subtitle, owner, system, desc in specs:
        W, H, inner = builder()
        p = SRC / folder / fn
        write_arch_svg(p, title, subtitle, W, H, inner, M(folder, owner, system))
        reg(p, folder, owner, system, title, desc)
    # roadmap (Gantt) — strategy
    folder = "03_strategy_initiatives"; W, H = 1300, 760
    inits = ["Kyriba Treasury Modernization", "SOX Controls Uplift", "Reporting Rationalization",
             "SI/AMS Vendor Optimization", "Snowflake Consolidation", "S/4HANA Migration Wave 2",
             "Enterprise AI Platform", "Cyber Zero-Trust"]
    swim = {"Treasury": PAL["ac_blue"], "Finance": PAL["ac_teal"], "Data": PAL["ac_purple"],
            "Procurement": PAL["ac_amber"], "IT": PAL["ac_green"], "Risk": PAL["ac_red"]}
    lanes = ["Treasury", "Risk", "Finance", "Procurement", "Data", "IT", "IT", "Risk"]
    qx = [180, 318, 456, 594, 732, 870, 1008, 1146]
    qlabels = ["Q1 FY26", "Q2 FY26", "Q3 FY26", "Q4 FY26", "Q1 FY27", "Q2 FY27", "Q3 FY27", "Q4 FY27"]
    inner = ""
    for i, q in enumerate(qlabels):
        inner += f'<line x1="{qx[i]}" y1="120" x2="{qx[i]}" y2="700" stroke="{PAL["line"]}" stroke-opacity="0.15"/>'
        inner += f'<text x="{qx[i]}" y="138" font-size="10.5" font-family="monospace" fill="{PAL["muted"]}">{q}</text>'
    milestones = []
    for i, n in enumerate(inits):
        y = 168 + i * 62
        start = random.randint(0, 3); span = random.randint(2, 5)
        x1 = qx[start]; x2 = qx[min(7, start + span)]
        c = swim[lanes[i]]
        inner += f'<text x="36" y="{y+22}" font-size="12" font-weight="600" fill="{PAL["ink"]}">{n}</text>'
        inner += f'<rect x="{x1}" y="{y}" width="{x2-x1}" height="30" rx="7" fill="{c}" opacity="0.9"/>'
        inner += f'<text x="{x1+10}" y="{y+20}" font-size="9.5" font-family="monospace" fill="#fff">{lanes[i]} · {span} quarters</text>'
        # milestone diamond at end
        inner += f'<path d="M{x2},{y-6} l8,8 l-8,8 l-8,-8 z" fill="{PAL["ink"]}"/>'
        milestones.append((x2, y))
    # dependency arrow examples
    inner += _edge(qx[2], 198, qx[2], 230, "dep", dashed=True)
    inner += _legend(1010, 600, [(swim["Treasury"], "Treasury"), (swim["Data"], "Data/Analytics"), (swim["Risk"], "Risk/Cyber")])
    p = SRC / folder / "roadmap_2026_2027.svg"
    write_arch_svg(p, "Transformation Roadmap FY26–FY27", "Swimlane roadmap with quarter grid, durations, and milestones", W, H, inner, M(folder, "Marcus Reilly", "Corporate Strategy"))
    reg(p, folder, "Marcus Reilly", "Corporate Strategy", "Transformation Roadmap FY26-27", "Swimlane roadmap with milestones and dependencies.")

# ============================================================================
# MAIN — build all + manifest + zip + dictionary + index
# ============================================================================
def build_root_artifacts():
    # README_LOAD_NOTES.md (zip root)
    readme = SRC.parent / "README_LOAD_NOTES.md"  # placeholder; real placement below
    # We place root deliverables under OUT and inside ZIP at root.

def main():
    for d in DOMAINS:
        (SRC / d).mkdir(parents=True, exist_ok=True)
    build_manifest_domain()
    build_enterprise_profile()
    build_org()
    build_strategy()
    build_finance()
    build_treasury()
    build_it()
    build_data()
    build_ops()
    build_servicenow()
    build_vendors()
    build_risk()
    build_ai()
    rich_diagrams()

    # ---- evidence_register.csv (ZIP root + OUT) ----
    ev_rows=[]
    for i, r in enumerate(REGISTRY):
        if r["evidence_usable_flag"]:
            ev_rows.append([f"EVID-{i+1:04d}", r["context_domain"], r["zip_path"], r["title"],
                            r["source_owner"], r["source_system"], r["source_date"], r["sensitivity"],
                            "synthetic", r["sha256"][:16]])
    ev_headers=["evidence_id","context_domain","source_path","title","source_owner","source_system",
                "source_date","sensitivity","provenance","sha256_prefix"]
    for target in [OUT / "evidence_register.csv", SRC.parent / "evidence_register.csv"]:
        pass
    ev_path = OUT / "evidence_register.csv"
    write_csv(ev_path, ev_headers, ev_rows, meta=M("00_manifest","AbarVa Delivery","AbarVa Generator"))

    # ---- data_dictionary.xlsx ----
    dd_path = OUT / "data_dictionary.xlsx"
    field_rows=[
        ["context_domain","One of 12 enterprise context domains","string"],
        ["source_owner","Accountable owner / function","string"],
        ["source_system","System of record","string"],
        ["source_date","As-of date (ISO)","date"],
        ["sensitivity","Internal | Confidential","enum"],
        ["synthetic_flag","Always true for this pack","bool"],
        ["evidence_usable_flag","Eligible as cited evidence","bool"],
        ["loader_route","Setup Admin loader path","string"],
        ["sha256","Content hash for integrity","string"],
        ["bytes","File size in bytes","int"],
    ]
    file_rows=[[r["zip_path"], r["file_type"], r["context_domain"], r["source_system"], r["bytes"]] for r in REGISTRY]
    write_xlsx(dd_path, {
        "Fields": (["field","description","type"], field_rows),
        "Files": (["zip_path","file_type","context_domain","source_system","bytes"], file_rows),
    }, title="Data Dictionary", meta=M("00_manifest","AbarVa Delivery","AbarVa Generator"))

    # ---- README_LOAD_NOTES.md ----
    readme_path = OUT / "README_LOAD_NOTES.md"
    by_domain={}
    for r in REGISTRY:
        by_domain.setdefault(r["context_domain"], 0)
        by_domain[r["context_domain"]]+=1
    write_md(readme_path, "README — Load Notes",
        textwrap.dedent(f"""
        ## {PACK_ID}
        Fully synthetic enterprise context pack for **{CLIENT_NAME}** (`{TENANT_KEY}`).

        - **Files:** {len(REGISTRY)} source documents (excludes root manifest/dictionary/register).
        - **Watermark:** every file carries `{WATERMARK_SHORT}` in content + metadata.
        - **manifest.json** at ZIP root carries per-file: context_domain, source_owner,
          source_system, source_date, sensitivity, synthetic_flag, evidence_usable_flag, loader_route.

        ## How to load (Setup Admin)
        1. Go to `/admin/context-layer/uploads` while switched to the Lakeshore tenant.
        2. Use the **Bulk / ZIP** connector; upload `{PACK_ID}.zip`.
        3. Choose mode `stage_and_process` (structured) / `stage_and_enqueue` (rich docs).
        4. Confirm attestation; the loader stages to Azure Blob `context-uploads/` and
           commits chunks to `enterprise_context_chunks`.
        5. Run `npm run embed:pending-chunks -- --tenant {TENANT_KEY}` (requires OPENAI_API_KEY)
           and the Azure AI Search backfill to make chunks searchable.

        ## Files per domain
        """) + "\n".join(f"- **{k}**: {v}" for k,v in sorted(by_domain.items())),
        meta=M("00_manifest","AbarVa Delivery","AbarVa Generator"))

    # ---- manifest.json ----
    manifest = {
        "pack_id": PACK_ID,
        "generated_date": GEN_DATE,
        "tenant_key": TENANT_KEY,
        "client_name": CLIENT_NAME,
        "synthetic": True,
        "synthetic_label": WATERMARK_SHORT,
        "watermark": WATERMARK,
        "loader": "setup-admin-bulk-zip",
        "manifest_version": "1.0",
        "file_count": len(REGISTRY),
        "domains": list(DOMAINS.values()),
        "root_artifacts": ["manifest.json","evidence_register.csv","data_dictionary.xlsx","README_LOAD_NOTES.md"],
        "files": [{
            "path": r["zip_path"],
            "name": r["file_name"],
            "file_type": r["file_type"],
            "bytes": r["bytes"],
            "sha256": r["sha256"],
            "title": r["title"],
            "description": r["description"],
            "context_domain": r["context_domain"],
            "source_owner": r["source_owner"],
            "source_system": r["source_system"],
            "source_date": r["source_date"],
            "sensitivity": r["sensitivity"],
            "synthetic_flag": r["synthetic_flag"],
            "evidence_usable_flag": r["evidence_usable_flag"],
            "loader_route": r["loader_route"],
        } for r in REGISTRY],
    }
    # place manifest in OUT and (copied) into zip root
    (OUT / f"{PACK_ID}_MANIFEST.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    # ---- ZIP (manifest.json + root artifacts + source/ tree) ----
    zip_path = OUT / f"{PACK_ID}.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("manifest.json", json.dumps(manifest, indent=2))
        z.write(readme_path, "README_LOAD_NOTES.md")
        z.write(ev_path, "evidence_register.csv")
        z.write(dd_path, "data_dictionary.xlsx")
        for r in REGISTRY:
            abs_path = OUT / r["file_path"]
            z.write(abs_path, r["zip_path"])
    zbytes = zip_path.stat().st_size

    # ---- GENERATED_FILE_INDEX.md ----
    idx_path = OUT / f"{PACK_ID}_GENERATED_FILE_INDEX.md"
    type_counts={}
    for r in REGISTRY:
        type_counts[r["file_type"]]=type_counts.get(r["file_type"],0)+1
    lines=[f"# {PACK_ID} — Generated File Index", "",
           f"- Generated: {GEN_DATE}", f"- Tenant: {CLIENT_NAME} (`{TENANT_KEY}`)",
           f"- Source files: **{len(REGISTRY)}**", f"- ZIP: `{PACK_ID}.zip` ({zbytes:,} bytes)",
           f"- Watermark: `{WATERMARK_SHORT}`", "",
           "## File type counts", "", "| type | count |", "|---|---|"]
    for t,c in sorted(type_counts.items()):
        lines.append(f"| {t} | {c} |")
    lines += ["", "## Files by domain", ""]
    cur=None
    for r in sorted(REGISTRY, key=lambda x: x["zip_path"]):
        dom=r["domain_folder"]
        if dom!=cur:
            lines += ["", f"### {dom} ({r['context_domain']})", "", "| file | type | bytes | owner | system | sensitivity | evidence |","|---|---|---|---|---|---|---|"]
            cur=dom
        lines.append(f"| `{r['file_name']}` | {r['file_type']} | {r['bytes']:,} | {r['source_owner']} | {r['source_system']} | {r['sensitivity']} | {'yes' if r['evidence_usable_flag'] else 'no'} |")
    idx_path.write_text("\n".join(lines)+"\n", encoding="utf-8")

    # summary json for downstream scripts
    (OUT / "parse-output" / "_generation_summary.json").write_text(json.dumps({
        "pack_id": PACK_ID, "file_count": len(REGISTRY), "zip_bytes": zbytes,
        "type_counts": type_counts,
        "total_source_bytes": sum(r["bytes"] for r in REGISTRY),
    }, indent=2), encoding="utf-8")

    print(f"\n=== DONE ===")
    print(f"files={len(REGISTRY)} zip_bytes={zbytes:,} types={type_counts}")
    print(f"source_bytes={sum(r['bytes'] for r in REGISTRY):,}")

if __name__ == "__main__":
    main()
