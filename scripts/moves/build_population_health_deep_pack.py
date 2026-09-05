#!/usr/bin/env python3
"""Build the synthetic population-health Moves rich-context fixture pack.

The pack is intentionally deterministic and unloaded. It gives testers a
non-trivial evidence set for the Moves context pack without turning fixture
content into a live tenant claim.
"""

from __future__ import annotations

import csv
import hashlib
import json
import shutil
import subprocess
import sys
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from textwrap import dedent
from typing import Iterable

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[2]
PACK_ROOT = ROOT / "docs/status/moves-rich-context/fixtures/population-health-command-center"
GENERATED_AT = "2026-09-05T00:00:00Z"
HEADLINE = {
    "attributed_lives": 418_000,
    "open_care_gaps": 1_142_000,
    "closure_rate_pct": 41.2,
    "weighted_cell_closure_rate_pct": 41.19,
    "quality_measures": 40,
}
PROHIBITED_STRINGS = [
    "New Mexico",
    "Albuquerque",
    "Santa Fe",
    "Las Cruces",
    "Rio Rancho",
    "Honolulu",
    "Hawaii",
    "Kona Coast",
    "Presbyterian",
]


@dataclass(frozen=True)
class TableSpec:
    slug: str
    label: str
    owner: str
    rows: list[dict[str, object]]


def stable_pick(values: list[str], index: int) -> str:
    return values[index % len(values)]


def allocate_total(total: int, count: int, seed: int, minimum: int = 1) -> list[int]:
    weights = [((i * 37 + seed * 17) % 97) + 20 for i in range(count)]
    remaining = total - (minimum * count)
    if remaining < 0:
        raise ValueError("total is below minimum allocation")
    weight_sum = sum(weights)
    values = [minimum + (remaining * weight // weight_sum) for weight in weights]
    delta = total - sum(values)
    order = sorted(range(count), key=lambda i: (-weights[i], i))
    for i in range(delta):
        values[order[i % count]] += 1
    return values


def write_csv(path: Path, rows: list[dict[str, object]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        raise ValueError(f"refusing to write empty csv: {path}")
    fieldnames = list(rows[0].keys())
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, lineterminator="\n")
        writer.writeheader()
        for row in rows:
            writer.writerow({key: row.get(key, "") for key in fieldnames})


def write_xlsx(path: Path, title: str, rows: list[dict[str, object]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        raise ValueError(f"refusing to write empty workbook: {path}")
    workbook = Workbook()
    workbook.properties.creator = "AbarVa synthetic fixture generator"
    workbook.properties.created = datetime(2026, 9, 5, tzinfo=timezone.utc)
    workbook.properties.modified = datetime(2026, 9, 5, tzinfo=timezone.utc)

    readme = workbook.active
    readme.title = "README"
    readme.append(["Purpose", "Synthetic Moves rich-context fixture. Not loaded into a tenant."])
    readme.append(["Generated", GENERATED_AT])
    readme.append(["Table", title])

    sheet = workbook.create_sheet("data")
    fieldnames = list(rows[0].keys())
    sheet.append(fieldnames)
    for row in rows:
        sheet.append([row.get(key, "") for key in fieldnames])

    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
    sheet.freeze_panes = "A2"
    for column in sheet.columns:
        width = min(max(len(str(cell.value or "")) for cell in column) + 2, 42)
        sheet.column_dimensions[get_column_letter(column[0].column)].width = width
    workbook.save(path)


def markdown_to_docx(markdown_path: Path, docx_path: Path) -> None:
    document = Document()
    document.core_properties.author = "AbarVa synthetic fixture generator"
    document.core_properties.created = datetime(2026, 9, 5, tzinfo=timezone.utc)
    document.core_properties.modified = datetime(2026, 9, 5, tzinfo=timezone.utc)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for raw in markdown_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def build_systems() -> list[dict[str, object]]:
    categories = ["EHR", "population health", "analytics", "claims", "care management", "integration", "finance", "security"]
    owners = ["VP Clinical Apps", "VP Value Based Care", "Director Data Platform", "Director Integration Services", "VP Application Services"]
    rows: list[dict[str, object]] = []
    special = [
        ("SYS-0001", "Enterprise EHR", "EHR", "Epic", "active", "daily", "yes", "owned"),
        ("SYS-0002", "Population Health Registry", "population health", "Epic Healthy Planet", "active", "daily", "yes", "owned"),
        ("SYS-0003", "Care Gap Outreach CRM", "care management", "Salesforce Health Cloud", "active", "daily", "yes", "owned"),
        ("SYS-0004", "Legacy Population Warehouse", "analytics", "Warehouse", "retired_2024", "none", "yes", "owned"),
        ("SYS-0005", "Sepsis Alert Sandbox", "clinical AI", "EHR alerting", "declined_2025", "none", "yes", "not_approved"),
        ("SYS-0006", "Island Legacy EHR", "EHR", "Oracle Health", "design_scope_only", "weekly", "yes", "named_gap"),
    ]
    for system_id, name, category, platform, lifecycle, cadence, phi, posture in special:
        rows.append(
            {
                "system_id": system_id,
                "system_name": name,
                "category": category,
                "owner_role": "VP Application Services" if system_id == "SYS-0006" else stable_pick(owners, len(rows)),
                "platform": platform,
                "hosting_model": stable_pick(["managed cloud", "private cloud", "hospital data center", "vendor hosted"], len(rows)),
                "lifecycle_state": lifecycle,
                "feed_cadence": cadence,
                "contains_phi_or_pii": phi,
                "integration_count": 12 + (len(rows) * 7) % 64,
                "monitoring_state": "not_monitored" if posture in {"not_approved", "named_gap"} else "monitored",
                "source_control_state": "catalogued",
                "evidence_note": posture,
            }
        )
    for index in range(7, 121):
        rows.append(
            {
                "system_id": f"SYS-{index:04d}",
                "system_name": f"Care delivery platform {index:03d}",
                "category": stable_pick(categories, index),
                "owner_role": stable_pick(owners, index),
                "platform": stable_pick(["Epic", "Databricks", "Snowflake", "ServiceNow", "Mirth", "Workday"], index),
                "hosting_model": stable_pick(["managed cloud", "private cloud", "hospital data center", "vendor hosted"], index),
                "lifecycle_state": stable_pick(["active", "active", "active", "rationalize", "monitor"], index),
                "feed_cadence": stable_pick(["daily", "daily", "near-real-time", "weekly", "monthly"], index),
                "contains_phi_or_pii": "yes" if index % 3 != 0 else "no",
                "integration_count": 4 + (index * 11) % 72,
                "monitoring_state": stable_pick(["monitored", "monitored", "partial", "not_monitored"], index),
                "source_control_state": stable_pick(["catalogued", "catalogued", "not_versioned"], index),
                "evidence_note": "synthetic system inventory row; no patient-level data",
            }
        )
    return rows


def build_interfaces() -> list[dict[str, object]]:
    rows = []
    protocols = ["HL7v2", "FHIR", "X12", "SFTP", "API", "flat-file"]
    for index in range(1, 87):
        not_versioned = index <= 41
        not_monitored = index <= 33
        rows.append(
            {
                "interface_id": f"INT-{index:04d}",
                "source_system_id": f"SYS-{((index * 5) % 120) + 1:04d}",
                "target_system_id": f"SYS-{((index * 11) % 120) + 1:04d}",
                "protocol": stable_pick(protocols, index),
                "channel_count": 1 + (index * 3) % 9,
                "source_control_state": "not_versioned" if not_versioned else "versioned",
                "monitoring_state": "not_monitored" if not_monitored else stable_pick(["monitored", "monitored", "partial"], index),
                "region_scope": stable_pick(["metro", "regional", "enterprise", "design_scope"], index),
                "phi_class": "phi" if index % 4 else "operations_only",
                "evidence_note": "synthetic interface inventory; 41 of 86 channels are intentionally unversioned" if not_versioned else "synthetic interface inventory",
            }
        )
    return rows


def build_facilities() -> list[dict[str, object]]:
    lives = allocate_total(418_000, 78, 9, minimum=850)
    rows = []
    for index, attributed in enumerate(lives, start=1):
        design_scope = "yes" if index in {7, 31, 64} else "no"
        rows.append(
            {
                "facility_id": f"FAC-{index:04d}",
                "facility_name": f"Care site {index:03d}",
                "region_label": stable_pick(["Metro North", "Metro South", "Central", "Coastal", "Mountain", "Virtual"], index),
                "facility_type": stable_pick(["hospital", "ambulatory", "clinic", "urgent care", "home health"], index),
                "ehr_platform": "Oracle Health" if design_scope == "yes" else "Epic",
                "feed_cadence": "weekly" if design_scope == "yes" else stable_pick(["daily", "daily", "near-real-time"], index),
                "design_scope_only": design_scope,
                "attributed_lives": attributed,
                "care_gap_capture_state": "design_only" if design_scope == "yes" else stable_pick(["captured", "captured", "partial"], index),
                "owner_role": stable_pick(["VP Clinical Operations", "Regional COO", "Director Ambulatory Ops"], index),
                "evidence_note": "legacy platform site held as design scope only" if design_scope == "yes" else "synthetic facility row",
            }
        )
    return rows


def build_org_roles() -> list[dict[str, object]]:
    rows = [
        {
            "role_id": "ROLE-0001",
            "role_title": "Chief Digital and Information Officer",
            "department": "Executive",
            "reports_to_role": "Chief Executive Officer",
            "location": "Metro HQ",
            "fte": 1.0,
            "vacancy_status": "filled",
            "primary_platform": "Enterprise",
            "decision_right": "executive sponsor",
            "evidence_note": "synthetic leadership role",
        },
        {
            "role_id": "ROLE-0002",
            "role_title": "VP Application Services",
            "department": "Technology",
            "reports_to_role": "Chief Digital and Information Officer",
            "location": "Metro HQ",
            "fte": 1.0,
            "vacancy_status": "open_requisition",
            "primary_platform": "EHR and clinical apps",
            "decision_right": "blocking vacancy",
            "evidence_note": "intentionally vacant; two directors report to this role",
        },
        {
            "role_id": "ROLE-0003",
            "role_title": "Director Clinical Applications",
            "department": "Technology",
            "reports_to_role": "VP Application Services",
            "location": "Regional Office",
            "fte": 1.0,
            "vacancy_status": "filled",
            "primary_platform": "EHR",
            "decision_right": "implementation owner",
            "evidence_note": "reports to vacant VP Application Services",
        },
        {
            "role_id": "ROLE-0004",
            "role_title": "Director Patient Access Platforms",
            "department": "Technology",
            "reports_to_role": "VP Application Services",
            "location": "Remote",
            "fte": 1.0,
            "vacancy_status": "filled",
            "primary_platform": "CRM",
            "decision_right": "workflow owner",
            "evidence_note": "reports to vacant VP Application Services",
        },
        {
            "role_id": "ROLE-0005",
            "role_title": "Data Governance Lead",
            "department": "Data and Analytics",
            "reports_to_role": "VP Enterprise Data",
            "location": "Metro HQ",
            "fte": 1.0,
            "vacancy_status": "open_requisition",
            "primary_platform": "Governance",
            "decision_right": "approval dependency",
            "evidence_note": "open role required before wave 1 governance sign-off",
        },
        {
            "role_id": "ROLE-0006",
            "role_title": "Clinical AI Monitoring Manager",
            "department": "Quality",
            "reports_to_role": "Chief Quality Officer",
            "location": "Remote",
            "fte": 1.0,
            "vacancy_status": "open_requisition",
            "primary_platform": "Model monitoring",
            "decision_right": "approval dependency",
            "evidence_note": "third open role; no model moves without owner",
        },
    ]
    departments = ["Quality", "Clinical Operations", "Value Based Care", "Data Platform", "Care Management", "Patient Access"]
    titles = ["Manager", "Analyst", "Nurse Informaticist", "Product Owner", "Data Engineer", "Quality Specialist"]
    for index in range(7, 73):
        rows.append(
            {
                "role_id": f"ROLE-{index:04d}",
                "role_title": f"{stable_pick(titles, index)} {index:03d}",
                "department": stable_pick(departments, index),
                "reports_to_role": stable_pick(["VP Value Based Care", "Chief Quality Officer", "Director Data Platform", "Director Clinical Applications"], index),
                "location": stable_pick(["Metro HQ", "Regional Office", "Remote"], index),
                "fte": 0.5 if index % 17 == 0 else 1.0,
                "vacancy_status": "filled",
                "primary_platform": stable_pick(["EHR", "CRM", "Analytics", "Care management"], index),
                "decision_right": stable_pick(["accountable", "consulted", "informed", "delivery"], index),
                "evidence_note": "synthetic org role row",
            }
        )
    return rows


def build_roadmap() -> list[dict[str, object]]:
    rows = []
    phases = ["P0 Originate", "P1 Charter", "P2 Discover", "P3 Design", "P4 Mobilize", "P5 Govern"]
    for index in range(1, 43):
        phase = stable_pick(phases, index - 1)
        rows.append(
            {
                "roadmap_id": f"RM-{index:04d}",
                "phase": phase,
                "wave": f"wave_{1 + ((index - 1) // 7)}",
                "initiative": stable_pick(["care gap closure", "quality measure reconciliation", "workflow redesign", "governance sequencing", "platform integration", "benefit proof"], index),
                "owner_role": stable_pick(["VP Value Based Care", "Chief Quality Officer", "VP Application Services", "Data Governance Lead"], index),
                "dependency": stable_pick(["data-quality gate", "clinical owner named", "interface monitoring", "benefit sign-off", "model monitoring plan"], index),
                "target_state": stable_pick(["decision-ready", "pilot-ready", "design-scope-only", "blocked pending owner"], index),
                "acceptance_signal": stable_pick(["coverage readout", "artifact cites evidence", "gate records carry-forward", "no invented value"], index),
                "evidence_note": "synthetic roadmap item",
            }
        )
    return rows


def build_quality_measures() -> list[dict[str, object]]:
    adult = [
        ("CBP", "Controlling High Blood Pressure"),
        ("HBD", "Hemoglobin A1c Control for Patients With Diabetes"),
        ("EED", "Eye Exam for Patients With Diabetes"),
        ("KED", "Kidney Health Evaluation for Patients With Diabetes"),
        ("COL-E", "Colorectal Cancer Screening"),
        ("BCS-E", "Breast Cancer Screening"),
        ("CCS-E", "Cervical Cancer Screening"),
        ("TRC", "Transitions of Care"),
        ("PCR", "Plan All-Cause Readmissions"),
        ("FUH", "Follow-Up After Hospitalization for Mental Illness"),
        ("FUM", "Follow-Up After Emergency Department Visit for Mental Illness"),
        ("FMC", "Follow-Up After ED Visit for People With Multiple High-Risk Chronic Conditions"),
        ("SNS-E", "Social Need Screening and Intervention"),
        ("AIS-E", "Adult Immunization Status"),
        ("TSC-E", "Tobacco Use Screening and Cessation Intervention"),
        ("PCE", "Pharmacotherapy Management of COPD Exacerbation"),
        ("OMW", "Osteoporosis Management in Women Who Had a Fracture"),
        ("SPC-E", "Statin Therapy for Patients With Cardiovascular Disease"),
        ("BPD-E", "Blood Pressure Control for Patients With Diabetes"),
    ]
    medicare = [
        ("MAD", "Medication Adherence for Diabetes Medications"),
        ("MAH", "Medication Adherence for Hypertension"),
        ("MAC", "Medication Adherence for Cholesterol"),
        ("SUPD", "Statin Use in Persons With Diabetes"),
        ("MRP", "Medication Reconciliation Post-Discharge"),
        ("COB", "Concurrent Use of Opioids and Benzodiazepines"),
        ("POLY-ACH", "Use of Multiple Anticholinergic Medications in Older Adults"),
        ("OMH", "Older Adult Health Outcomes Review"),
    ]
    child = [
        ("CIS-E", "Childhood Immunization Status"),
        ("IMA-E", "Immunizations for Adolescents"),
        ("W30", "Well-Child Visits in the First 30 Months of Life"),
        ("WCV", "Child and Adolescent Well-Care Visits"),
        ("LSC-E", "Lead Screening in Children"),
        ("ADD-E", "Follow-Up Care for Children Prescribed ADHD Medication"),
        ("CHL", "Chlamydia Screening in Women"),
        ("AAB", "Avoidance of Antibiotic Treatment for Acute Bronchitis"),
    ]
    women = [
        ("PPC", "Prenatal and Postpartum Care"),
        ("FUA", "Follow-Up After Emergency Department Visit for Substance Use"),
        ("AMM", "Antidepressant Medication Management"),
        ("GSD", "Glycemic Status Assessment for Patients With Diabetes"),
        ("SAA", "Adherence to Antipsychotic Medications for Individuals With Schizophrenia"),
    ]
    specs = [("adult", adult), ("medicare", medicare), ("child", child), ("women", women)]
    rows = []
    index = 0
    for category, measures in specs:
        for code, name in measures:
            index += 1
            rows.append(
                {
                    "measure_id": f"QM-{index:04d}",
                    "measure_code": code,
                    "measure_name": name,
                    "domain": category,
                    "stars_relevant": "yes" if category in {"adult", "medicare"} else stable_pick(["yes", "no"], index),
                    "cohort_applicability": {
                        "adult": "ma_seniors|medicare_snp|medicaid_adult|commercial_adult",
                        "medicare": "ma_seniors|medicare_snp",
                        "child": "pediatric_medicaid|pediatric_commercial",
                        "women": "medicaid_adult|commercial_adult",
                    }[category],
                    "source_reference": "NCQA HEDIS technical resources and CMS Star Ratings measure categories",
                    "evidence_note": "synthetic measure catalog; aggregate only",
                }
            )
    if len(rows) != 40:
        raise AssertionError(f"expected 40 quality measures, got {len(rows)}")
    return rows


def build_care_gap_cells(measures: list[dict[str, object]]) -> list[dict[str, object]]:
    cohorts = {
        "ma_seniors": ("Medicare Advantage seniors", 119_000),
        "medicare_snp": ("Medicare SNP", 38_000),
        "medicaid_adult": ("Medicaid adult", 96_000),
        "commercial_adult": ("Commercial adult", 107_000),
        "pediatric_medicaid": ("Pediatric Medicaid", 36_000),
        "pediatric_commercial": ("Pediatric commercial", 22_000),
    }
    measure_rows: list[tuple[dict[str, object], str]] = []
    for measure in measures:
        for cohort_key in str(measure["cohort_applicability"]).split("|"):
            measure_rows.append((measure, cohort_key))
    if len(measure_rows) != 118:
        raise AssertionError(f"expected 118 measure/cohort cells, got {len(measure_rows)}")
    opens = allocate_total(HEADLINE["open_care_gaps"], len(measure_rows), 41, minimum=275)
    target_closed = round(
        HEADLINE["open_care_gaps"]
        * HEADLINE["weighted_cell_closure_rate_pct"]
        / (100 - HEADLINE["weighted_cell_closure_rate_pct"])
    )
    closed = allocate_total(target_closed, len(measure_rows), 73, minimum=160)
    rows = []
    for index, ((measure, cohort_key), open_count, closed_count) in enumerate(zip(measure_rows, opens, closed), start=1):
        cohort_label, lives = cohorts[cohort_key]
        denominator = open_count + closed_count
        rows.append(
            {
                "cell_id": f"CG-{index:04d}",
                "measure_code": measure["measure_code"],
                "measure_name": measure["measure_name"],
                "cohort_key": cohort_key,
                "cohort_label": cohort_label,
                "cohort_attributed_lives": lives,
                "open_gap_count": open_count,
                "closed_gap_count": closed_count,
                "denominator_count": denominator,
                "closure_rate_pct": round((closed_count / denominator) * 100, 2),
                "line_of_business": stable_pick(["Medicare", "Medicaid", "Commercial"], index),
                "source_system": stable_pick(["EHR quality extract", "plan supplemental file", "care management registry"], index),
                "evidence_note": "synthetic aggregate measure/cohort cell; no member-level data",
            }
        )
    return rows


def build_operating_evidence() -> list[dict[str, object]]:
    rows = []

    def add(record_type: str, count: int, label: str, severity: str = "medium") -> None:
        start = len(rows) + 1
        for offset in range(count):
            index = start + offset
            rows.append(
                {
                    "record_id": f"OE-{index:04d}",
                    "record_type": record_type,
                    "label": f"{label} {offset + 1:03d}",
                    "owner_role": stable_pick(["VP Value Based Care", "Chief Quality Officer", "Director Data Platform", "VP Application Services", "Clinical Operations Lead"], index),
                    "status": stable_pick(["open", "validated", "partial", "blocked", "design_scope"], index),
                    "severity": severity if offset % 5 == 0 else stable_pick(["low", "medium", "high"], index),
                    "metric_value": (index * 13) % 101,
                    "source_reference": stable_pick(["workshop note", "extract inventory", "governance log", "process observation"], index),
                    "evidence_note": "synthetic operating evidence row",
                }
            )

    add("site", 48, "Care delivery site")
    add("stakeholder", 34, "Named stakeholder role")
    add("data_quality_finding", 30, "Data quality finding", "high")
    add("risk", 28, "Delivery risk", "high")
    for minute in range(1, 54):
        rows.append(
            {
                "record_id": f"OE-{len(rows) + 1:04d}",
                "record_type": "time_motion_minute",
                "label": f"Care manager hour minute {minute:02d}",
                "owner_role": "Care Management Lead",
                "status": "search_or_reconciliation" if minute <= 31 else "clinical_action",
                "severity": "medium",
                "metric_value": 1,
                "source_reference": "synthetic time-and-motion table",
                "evidence_note": "31 of 53 observed minutes intentionally assigned to search and reconciliation",
            }
        )
    add("evidence_artifact", 22, "Uploadable evidence artifact")
    add("constraint", 6, "Delivery constraint")
    add("phase_acceptance_check", 21, "Phase acceptance check")
    add("prior_attempt", 3, "Prior failed attempt", "high")
    for value_name in ["actuarial shared-savings model", "medical expense attribution", "staff productivity conversion"]:
        rows.append(
            {
                "record_id": f"OE-{len(rows) + 1:04d}",
                "record_type": "value_input",
                "label": value_name,
                "owner_role": stable_pick(["Actuary", "Chief Financial Officer", "Clinical Operations Lead"], len(rows)),
                "status": "UNVALIDATED",
                "severity": "high",
                "metric_value": 0,
                "source_reference": "pending finance validation",
                "evidence_note": "zero value until owner sign-off; do not use as realized value",
            }
        )
    add("value_exclusion", 3, "Named value exclusion")
    add("upload_block", 5, "Upload block")
    add("runbook_step", 26, "Tester runbook step")
    if len(rows) != 282:
        raise AssertionError(f"expected 282 operating evidence rows, got {len(rows)}")
    return rows


def build_tables() -> list[TableSpec]:
    quality = build_quality_measures()
    return [
        TableSpec("systems_inventory", "Systems Inventory", "IT applications owner", build_systems()),
        TableSpec("interface_inventory", "Interface Inventory", "Integration services owner", build_interfaces()),
        TableSpec("facility_footprint", "Facility Footprint", "Clinical operations owner", build_facilities()),
        TableSpec("org_roles", "Organization Roles", "Technology leadership", build_org_roles()),
        TableSpec("roadmap", "Roadmap", "Move sponsor", build_roadmap()),
        TableSpec("quality_measures", "Quality Measures", "Quality analytics", quality),
        TableSpec("care_gap_cells", "Care Gap Cells", "Population health analytics", build_care_gap_cells(quality)),
        TableSpec("operating_evidence", "Operating Evidence", "Move PMO", build_operating_evidence()),
    ]


def write_markdown_documents(doc_dir: Path) -> list[Path]:
    docs = {
        "case_study.md": f"""# Population Health Command Center Case Study

## 1. Organization
Healthcare Demo Network is a synthetic integrated delivery and health-plan enterprise with {HEADLINE['attributed_lives']:,} attributed lives. The pack uses cover names, aggregate counts, and no patient-level data.

## 2. Problem Mechanics
The current-state evidence shows {HEADLINE['open_care_gaps']:,} open care gaps across quality, access, chronic-condition, pediatric, and medication-adherence measures. Provider and plan measures intentionally disagree in roughly 29 percent of operating records so the Move has to reconcile source authority rather than summarize one tidy table.

## 3. Why Previous Attempts Failed
A readmission model was retired in 2024 after workflow adoption and monitoring fell behind the stated model performance. A sepsis alert was declined in 2025 because the submission had no monitoring plan and no named clinical owner. A care-gap dashboard never cleared finance review because value inputs were treated as benefits before actuarial sign-off.

## 4. Constraints
The design is constrained by source disagreement, a weekly legacy-platform feed for the island market, a vacant VP Application Services role, missing interface monitoring, unvalidated value inputs, and a required clinical AI governance owner before wave 1.

## 5. Decision Required
The Move must decide whether to charter a governed command center, sequence the first wave around care-gap closure and source reconciliation, and defer model-assist features until monitoring ownership is named.

## 6. Phase Expectations
P0 must name the problem and rejected claims. P1 must identify owners and value inputs. P2 must surface coverage and conflict. P3 must design the data and workflow route. P4 must mobilize monitoring and adoption. P5 must govern benefit readback without converting unvalidated estimates into realized value.

## 7. Moment Worth Showing
The artifact should cite the shadow care-gap registry, the unversioned interface channels, the retired 2024 readmission model, the declined 2025 sepsis alert, and the zero-valued UNVALIDATED benefit inputs.

## 8. What This Does Not Claim
This pack is not loaded into any tenant. It is not a real client record, not a PHI dataset, not a CMS submission, not an actuarial model, and not proof that the live prompt packed the evidence.
""",
        "process_walkthrough.md": """# Population Health Workflow Walkthrough

## Operating Friction
The care-management time-and-motion evidence assigns 31 of every 53 minutes to search and reconciliation before clinical action. That makes the caseload argument arithmetic: reducing search time is the only way the same team can close more gaps without pretending staffing changed.

## Data Path
Structured extracts land through the current-state path as CSV or XLSX. Narrative context lands through Upload and Review as Markdown or DOCX. Both paths carry the same content so testers can compare parse coverage and prompt coverage honestly.

## Reconciliation Logic
Provider-sourced measure status is treated as workflow truth. Plan-sourced status is treated as payment and Stars truth. Conflicts do not fail the Move; they become named evidence requests and phase carry-forward items.

## Design Scope Boundary
The island legacy platform remains design-scope only because its feed is weekly and its quality-measure extract is incomplete. The Move may discuss sequencing, but it may not count that market as ready for wave 1 automation.
""",
        "governance_history.md": """# Governance History

## Retired Readmission Model
The 2024 readmission model is included to catch shallow synthesis. A correct artifact should say the model was retired and should not present it as an available accelerator.

## Declined Sepsis Alert
The 2025 sepsis alert was declined by the AI Governance Council because the operating submission lacked a monitoring plan and a named clinical owner. The refusal was not an accuracy finding, so a correct artifact should not describe it as model underperformance.

## Benefit Discipline
Three value inputs remain zero and UNVALIDATED until the actuary and finance owner sign them off. This is deliberate; a generated business case should treat them as open inputs, not benefits.
""",
        "test_runbook.md": """# Tester Runbook

## Upload Blocks
1. Upload the eight CSV files through the structured current-state path.
2. Upload the eight XLSX files through the same structured path and compare row counts.
3. Upload the four DOCX files through Upload and Review.
4. Upload selected Markdown companions through Upload and Review.
5. Capture context coverage and generated-artifact citations after P3 generation.

## Required Signals
The generated artifact should mention the shadow registry, unversioned interface channels, weekly legacy-platform feed, retired readmission model, declined sepsis alert, vacant VP Application Services role, and UNVALIDATED zero-value inputs.

## Failure Modes
Coverage telemetry can be green while the artifact ignores the high-signal evidence. This pack gives testers concrete needles to look for so prompt-packing and synthesis can be evaluated separately.
""",
    }
    paths = []
    for filename, content in docs.items():
        path = doc_dir / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(dedent(content).strip() + "\n", encoding="utf-8")
        paths.append(path)
    for md_path in paths:
        markdown_to_docx(md_path, md_path.with_suffix(".docx"))
    return paths


def write_cheat_sheet(path: Path, file_refs: list[str]) -> None:
    phases = ["P0", "P1", "P2", "P3", "P4", "P5"]
    fields = [
        "tenant_key",
        "source_file",
        "source_row",
        "owner_role",
        "evidence_note",
        "measure_code",
        "cohort_key",
        "open_gap_count",
        "closed_gap_count",
        "closure_rate_pct",
        "source_control_state",
        "monitoring_state",
        "vacancy_status",
        "feed_cadence",
        "design_scope_only",
        "decision_right",
        "status",
        "severity",
        "metric_value",
        "source_reference",
        "acceptance_signal",
    ]
    blocks = [
        ("CSV current-state", "Structured path", "Use all eight CSV tables."),
        ("XLSX current-state", "Structured path", "Use all eight XLSX workbooks."),
        ("DOCX narrative", "Upload and Review", "Use all four DOCX narratives."),
        ("Markdown narrative", "Upload and Review", "Use selected Markdown companions."),
        ("Control proof", "Reviewer", "Use expected signals and validation report."),
    ]
    ref_items = "\n".join(f"<li><code>{ref}</code></li>" for ref in file_refs)
    phase_items = "\n".join(f"<span class=\"pill\" data-phase=\"{phase}\">{phase}</span>" for phase in phases)
    field_items = "\n".join(f"<li data-field=\"{field}\">{field}</li>" for field in fields)
    block_items = "\n".join(
        f"<article class=\"block\" data-upload-block=\"{name}\"><h3>{name}</h3><b>{pathway}</b><p>{copy}</p></article>"
        for name, pathway, copy in blocks
    )
    html = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>Population Health Command Center Fixture Cheat Sheet</title>
  <style>
    :root {{ color-scheme: light; --ink:#17202a; --muted:#51606f; --line:#c9d3dd; --soft:#f4f7fa; --accent:#0f766e; --warn:#9a3412; }}
    body {{ margin:0; font-family: Arial, sans-serif; color:var(--ink); background:#ffffff; }}
    main {{ max-width:1180px; margin:0 auto; padding:28px; }}
    h1 {{ font-size:32px; margin:0 0 8px; }}
    h2 {{ font-size:18px; margin:28px 0 10px; }}
    p {{ color:var(--muted); line-height:1.45; }}
    .grid {{ display:grid; grid-template-columns:repeat(5, minmax(150px, 1fr)); gap:10px; }}
    .block {{ border:1px solid var(--line); border-radius:6px; padding:12px; background:var(--soft); min-height:118px; }}
    .block h3 {{ margin:0 0 8px; font-size:15px; }}
    .phases {{ display:flex; gap:8px; flex-wrap:wrap; }}
    .pill {{ border:1px solid var(--accent); color:var(--accent); padding:7px 10px; border-radius:999px; font-weight:700; }}
    .fields {{ columns:3; border:1px solid var(--line); border-radius:6px; padding:14px 24px; }}
    .refs {{ columns:2; border:1px solid var(--line); border-radius:6px; padding:14px 24px; }}
    code {{ overflow-wrap:anywhere; }}
    .note {{ border-left:4px solid var(--warn); padding-left:12px; }}
    @media (max-width: 760px) {{ main {{ padding:18px; }} .grid {{ grid-template-columns:1fr; }} .fields, .refs {{ columns:1; }} }}
  </style>
</head>
<body>
<main data-fixture="population-health-command-center" data-phase-count="6" data-field-count="21" data-upload-block-count="5" data-file-ref-count="{len(file_refs)}">
  <h1>Population Health Command Center Fixture</h1>
  <p>Unloaded synthetic evidence pack for Moves rich-context testing. It is aggregate-only, uses cover names, and does not claim live tenant readiness.</p>
  <section>
    <h2>Phases</h2>
    <div class="phases">{phase_items}</div>
  </section>
  <section>
    <h2>Upload Blocks</h2>
    <div class="grid">{block_items}</div>
  </section>
  <section>
    <h2>Fields To Watch</h2>
    <ol class="fields">{field_items}</ol>
  </section>
  <section>
    <h2>File References</h2>
    <ol class="refs" data-file-refs>{ref_items}</ol>
  </section>
  <p class="note">Expected generated-artifact signals: shadow registry, unversioned channels, weekly legacy feed, retired model, declined alert, vacant VP Application Services, and zero UNVALIDATED value inputs.</p>
</main>
</body>
</html>
"""
    path.write_text(html, encoding="utf-8")


def write_load_plans(load_dir: Path) -> list[Path]:
    plans = {
        "01-structured-csv-upload.json": {"path": "structured_current_state_csv", "files": 8, "auto_commit_expected": True},
        "02-structured-xlsx-upload.json": {"path": "structured_current_state_xlsx", "files": 8, "auto_commit_expected": True},
        "03-review-docx-upload.json": {"path": "upload_and_review_docx", "files": 4, "auto_commit_expected": False},
        "04-review-markdown-upload.json": {"path": "upload_and_review_markdown", "files": 4, "auto_commit_expected": False},
    }
    paths = []
    load_dir.mkdir(parents=True, exist_ok=True)
    for name, payload in plans.items():
        path = load_dir / name
        path.write_text(json.dumps({"generated_at": GENERATED_AT, **payload}, indent=2) + "\n", encoding="utf-8")
        paths.append(path)
    return paths


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_index_and_checksums(pack_files: Iterable[Path]) -> None:
    excluded = {"package-index.csv", "checksums.sha256"}
    files = sorted(
        {
            path
            for path in pack_files
            if path.is_file() and path.relative_to(PACK_ROOT).as_posix() not in excluded
        }
    )
    with (PACK_ROOT / "package-index.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=["relative_path", "bytes", "sha256"],
            lineterminator="\n",
        )
        writer.writeheader()
        for path in files:
            writer.writerow(
                {
                    "relative_path": path.relative_to(PACK_ROOT).as_posix(),
                    "bytes": path.stat().st_size,
                    "sha256": sha256_file(path),
                }
            )
    lines = [f"{sha256_file(path)}  {path.relative_to(PACK_ROOT).as_posix()}" for path in files]
    (PACK_ROOT / "checksums.sha256").write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_pack() -> None:
    if PACK_ROOT.exists():
        shutil.rmtree(PACK_ROOT)
    data_dir = PACK_ROOT / "data"
    doc_dir = PACK_ROOT / "documents"
    load_dir = PACK_ROOT / "load-plans"
    qa_dir = PACK_ROOT / "qa"
    qa_dir.mkdir(parents=True, exist_ok=True)

    tables = build_tables()
    csv_paths: list[Path] = []
    xlsx_paths: list[Path] = []
    for table in tables:
        csv_path = data_dir / f"{table.slug}.csv"
        xlsx_path = data_dir / f"{table.slug}.xlsx"
        write_csv(csv_path, table.rows)
        write_xlsx(xlsx_path, table.label, table.rows)
        csv_paths.append(csv_path)
        xlsx_paths.append(xlsx_path)

    md_paths = write_markdown_documents(doc_dir)
    docx_paths = [path.with_suffix(".docx") for path in md_paths]
    load_paths = write_load_plans(load_dir)

    expected_signals = {
        "generated_at": GENERATED_AT,
        "must_mention": [
            "shadow registry",
            "unversioned interface channels",
            "weekly legacy-platform feed",
            "retired 2024 readmission model",
            "declined 2025 sepsis alert",
            "vacant VP Application Services",
            "UNVALIDATED zero-value inputs",
        ],
        "must_not_claim": [
            "loaded into tenant",
            "live prompt proof",
            "real patient data",
            "realized savings",
        ],
    }
    (PACK_ROOT / "expected-signals.json").write_text(json.dumps(expected_signals, indent=2) + "\n", encoding="utf-8")
    (qa_dir / "deep-pack-contract.json").write_text(
        json.dumps(
            {
                "generated_at": GENERATED_AT,
                "expected_files_after_validation": 38,
                "expected_structured_tables": 8,
                "expected_data_rows": 838,
                "expected_cheat_sheet_refs": 23,
                "validation_scope": [
                    "csv_xlsx_row_parity",
                    "docx_markdown_heading_parity",
                    "care_gap_reconciliation",
                    "cheat_sheet_references",
                    "public_repo_sensitivity_scan",
                ],
            },
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )

    acceptance_rows = [
        {"phase": "P0", "checks": 3, "evidence_signal": "problem and non-claims separated"},
        {"phase": "P1", "checks": 4, "evidence_signal": "owners, value inputs, vacancies named"},
        {"phase": "P2", "checks": 4, "evidence_signal": "coverage, conflicts, and missing feeds surfaced"},
        {"phase": "P3", "checks": 4, "evidence_signal": "data/workflow/governance sequence designed"},
        {"phase": "P4", "checks": 3, "evidence_signal": "monitoring and adoption plan mobilized"},
        {"phase": "P5", "checks": 3, "evidence_signal": "benefit readback holds zero until sign-off"},
    ]
    write_csv(PACK_ROOT / "acceptance-matrix.csv", acceptance_rows)

    cheat_refs = [
        *(path.relative_to(PACK_ROOT).as_posix() for path in csv_paths),
        *(path.relative_to(PACK_ROOT).as_posix() for path in xlsx_paths),
        *(path.relative_to(PACK_ROOT).as_posix() for path in docx_paths),
        "documents/case_study.md",
        "documents/process_walkthrough.md",
        "expected-signals.json",
    ]
    if len(cheat_refs) != 23:
        raise AssertionError(f"expected 23 cheat-sheet refs, got {len(cheat_refs)}")
    write_cheat_sheet(PACK_ROOT / "upload-cheat-sheet.html", cheat_refs)

    manifest = {
        "pack_id": "population-health-command-center-rich-context-v1",
        "generated_at": GENERATED_AT,
        "fixture_status": "synthetic_unloaded",
        "tenant_loading_status": "not_loaded",
        "public_repo_safety": "uses cover organization names, generic locations, aggregate counts, and no PHI/PII",
        "headline": HEADLINE,
        "tables": [
            {
                "slug": table.slug,
                "label": table.label,
                "owner": table.owner,
                "rows": len(table.rows),
                "csv": (data_dir / f"{table.slug}.csv").relative_to(PACK_ROOT).as_posix(),
                "xlsx": (data_dir / f"{table.slug}.xlsx").relative_to(PACK_ROOT).as_posix(),
            }
            for table in tables
        ],
        "narrative_documents": [
            {
                "markdown": path.relative_to(PACK_ROOT).as_posix(),
                "docx": path.with_suffix(".docx").relative_to(PACK_ROOT).as_posix(),
            }
            for path in md_paths
        ],
        "upload_blocks": [path.relative_to(PACK_ROOT).as_posix() for path in load_paths],
        "cheat_sheet_file_refs": cheat_refs,
        "expected_artifact_signals": expected_signals["must_mention"],
        "non_claims": expected_signals["must_not_claim"],
    }
    (PACK_ROOT / "manifest.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")

    readme = f"""# Population Health Command Center Rich-Context Fixture

This is an unloaded synthetic evidence pack for exercising Moves rich-context ingestion and prompt coverage. It is not a client record, not PHI, and not loaded into any tenant.

## Contents

- 8 structured tables, each as CSV and XLSX.
- 4 narrative documents, each as Markdown and DOCX.
- 838 CSV data rows across the structured tables.
- 118 measure-by-cohort care-gap cells that reconcile to {HEADLINE['open_care_gaps']:,} open gaps.
- 40 quality measures, including Medicare-only Part D/Stars measures and pediatric-only measures.
- A cheat sheet with 6 phases, 21 watched fields, 5 upload blocks, and 23 file references.

## Use

Run `npm run moves:rich-context-pack:validate` after regeneration. The validator re-opens the XLSX and DOCX files, reconciles care-gap totals, checks upload references, and scans for prohibited location/client hints.

## Boundary

This pack is fixture evidence only. Loading, approval, prompt packing, and generated-artifact citation proof are separate test steps.
"""
    (PACK_ROOT / "README.md").write_text(readme, encoding="utf-8")

    all_files = [path for path in PACK_ROOT.rglob("*") if path.is_file()]
    write_index_and_checksums(all_files)


def run_validator() -> None:
    subprocess.run([sys.executable, str(ROOT / "scripts/moves/validate_population_health_deep_pack.py")], cwd=ROOT, check=True)


def main() -> None:
    build_pack()
    if "--validate" in sys.argv:
        run_validator()
        write_index_and_checksums(path for path in PACK_ROOT.rglob("*") if path.is_file())
    print(f"built {PACK_ROOT}")


if __name__ == "__main__":
    main()
