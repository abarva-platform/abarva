#!/usr/bin/env python3
"""Validate the synthetic population-health Moves rich-context fixture pack."""

from __future__ import annotations

import csv
import html.parser
import json
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[2]
PACK_ROOT = ROOT / "docs/status/moves-rich-context/fixtures/population-health-command-center"
REPORT_JSON = PACK_ROOT / "validation-report.json"
REPORT_MD = PACK_ROOT / "validation-report.md"
EXPECTED = {
    "file_count": 38,
    "data_rows": 838,
    "table_count": 8,
    "narrative_count": 4,
    "care_gap_cells": 118,
    "open_care_gaps": 1_142_000,
    "attributed_lives": 418_000,
    "quality_measures": 40,
    "weighted_cell_closure_rate_pct": 41.19,
    "unversioned_interfaces": 41,
    "unmonitored_interfaces": 33,
    "org_open_requisitions": 3,
    "cheat_sheet_phases": 6,
    "cheat_sheet_fields": 21,
    "cheat_sheet_upload_blocks": 5,
    "cheat_sheet_file_refs": 23,
}
PROHIBITED_PATTERNS = [
    r"\bNew Mexico\b",
    r"\bAlbuquerque\b",
    r"\bSanta Fe\b",
    r"\bLas Cruces\b",
    r"\bRio Rancho\b",
    r"\bHonolulu\b",
    r"\bHawaii\b",
    r"\bKona Coast\b",
    r"\bPresbyterian\b",
    r"\bPHS\b",
]


class CheatSheetParser(html.parser.HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.phase_count = 0
        self.field_count = 0
        self.upload_block_count = 0
        self.file_refs: list[str] = []
        self._in_ref_code = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attr = dict(attrs)
        if attr.get("data-phase"):
            self.phase_count += 1
        if attr.get("data-field"):
            self.field_count += 1
        if attr.get("data-upload-block"):
            self.upload_block_count += 1
        if tag == "code":
            self._in_ref_code = True

    def handle_endtag(self, tag: str) -> None:
        if tag == "code":
            self._in_ref_code = False

    def handle_data(self, data: str) -> None:
        if self._in_ref_code:
            value = data.strip()
            if "/" in value or value.endswith(".json"):
                self.file_refs.append(value)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def normalize(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def equivalent_cell(left: object, right: object) -> bool:
    left_value = normalize(left)
    right_value = normalize(right)
    if left_value == right_value:
        return True
    try:
        return float(left_value) == float(right_value)
    except ValueError:
        return False


def read_xlsx_data(path: Path) -> list[dict[str, str]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    if "data" not in workbook.sheetnames:
        raise AssertionError(f"{path} has no data sheet")
    sheet = workbook["data"]
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        raise AssertionError(f"{path} data sheet is empty")
    headers = [normalize(cell) for cell in rows[0]]
    return [{headers[index]: normalize(value) for index, value in enumerate(row)} for row in rows[1:]]


def numeric(row: dict[str, str], key: str) -> int:
    return int(float(row[key]))


def fail(failures: list[str], message: str) -> None:
    failures.append(message)


def compare_rows(csv_rows: list[dict[str, str]], xlsx_rows: list[dict[str, str]], label: str, failures: list[str]) -> None:
    if len(csv_rows) != len(xlsx_rows):
        fail(failures, f"{label}: CSV/XLSX row count mismatch {len(csv_rows)} != {len(xlsx_rows)}")
        return
    csv_headers = list(csv_rows[0].keys()) if csv_rows else []
    xlsx_headers = list(xlsx_rows[0].keys()) if xlsx_rows else []
    if csv_headers != xlsx_headers:
        fail(failures, f"{label}: CSV/XLSX headers mismatch")
        return
    for index, (csv_row, xlsx_row) in enumerate(zip(csv_rows, xlsx_rows), start=1):
        for header in csv_headers:
            if not equivalent_cell(csv_row[header], xlsx_row[header]):
                fail(failures, f"{label}: row {index} column {header} mismatch")
                return


def validate_doc_pair(markdown_path: Path, docx_path: Path, failures: list[str]) -> None:
    if not docx_path.exists():
        fail(failures, f"missing DOCX companion for {markdown_path.relative_to(PACK_ROOT)}")
        return
    markdown = markdown_path.read_text(encoding="utf-8")
    document = Document(docx_path)
    docx_text = "\n".join(p.text for p in document.paragraphs)
    headings = [line.lstrip("# ").strip() for line in markdown.splitlines() if line.startswith("#")]
    for heading in headings:
        if heading and heading not in docx_text:
            fail(failures, f"{docx_path.name}: missing heading {heading!r}")
    for phrase in ["synthetic", "not loaded", "UNVALIDATED"]:
        if phrase in markdown and phrase not in docx_text:
            fail(failures, f"{docx_path.name}: missing key phrase {phrase!r}")


def validate_cheat_sheet(failures: list[str]) -> dict[str, object]:
    parser = CheatSheetParser()
    parser.feed((PACK_ROOT / "upload-cheat-sheet.html").read_text(encoding="utf-8"))
    if parser.phase_count != EXPECTED["cheat_sheet_phases"]:
        fail(failures, f"cheat sheet phase count {parser.phase_count} != {EXPECTED['cheat_sheet_phases']}")
    if parser.field_count != EXPECTED["cheat_sheet_fields"]:
        fail(failures, f"cheat sheet field count {parser.field_count} != {EXPECTED['cheat_sheet_fields']}")
    if parser.upload_block_count != EXPECTED["cheat_sheet_upload_blocks"]:
        fail(failures, f"cheat sheet upload block count {parser.upload_block_count} != {EXPECTED['cheat_sheet_upload_blocks']}")
    if len(parser.file_refs) != EXPECTED["cheat_sheet_file_refs"]:
        fail(failures, f"cheat sheet file refs {len(parser.file_refs)} != {EXPECTED['cheat_sheet_file_refs']}")
    missing = [ref for ref in parser.file_refs if not (PACK_ROOT / ref).exists()]
    if missing:
        fail(failures, f"cheat sheet missing refs: {missing}")
    return {
        "phases": parser.phase_count,
        "fields": parser.field_count,
        "upload_blocks": parser.upload_block_count,
        "file_refs": len(parser.file_refs),
        "missing_refs": missing,
    }


def scan_prohibited_strings(failures: list[str]) -> list[dict[str, str]]:
    matches: list[dict[str, str]] = []
    paths = [path for path in PACK_ROOT.rglob("*") if path.is_file()]

    def extract_text(path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".csv", ".md", ".html", ".json", ".txt", ".sha256"}:
            return path.read_text(encoding="utf-8")
        if suffix == ".docx":
            document = Document(path)
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        if suffix == ".xlsx":
            workbook = load_workbook(path, read_only=True, data_only=True)
            values: list[str] = []
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    values.extend(normalize(cell) for cell in row if cell is not None)
            return "\n".join(values)
        return ""

    for path in paths:
        text = extract_text(path)
        for pattern in PROHIBITED_PATTERNS:
            if re.search(pattern, text):
                matches.append({"file": path.relative_to(PACK_ROOT).as_posix(), "pattern": pattern})
    if matches:
        fail(failures, f"prohibited strings present: {matches}")
    return matches


def validate_pack() -> dict[str, object]:
    failures: list[str] = []
    if not PACK_ROOT.exists():
        raise SystemExit(f"missing pack root: {PACK_ROOT}")

    manifest = json.loads((PACK_ROOT / "manifest.json").read_text(encoding="utf-8"))
    tables = manifest["tables"]
    if len(tables) != EXPECTED["table_count"]:
        fail(failures, f"table count {len(tables)} != {EXPECTED['table_count']}")
    if len(manifest["narrative_documents"]) != EXPECTED["narrative_count"]:
        fail(failures, f"narrative count {len(manifest['narrative_documents'])} != {EXPECTED['narrative_count']}")

    table_rows: dict[str, list[dict[str, str]]] = {}
    total_rows = 0
    for table in tables:
        csv_path = PACK_ROOT / table["csv"]
        xlsx_path = PACK_ROOT / table["xlsx"]
        rows = read_csv(csv_path)
        xlsx_rows = read_xlsx_data(xlsx_path)
        compare_rows(rows, xlsx_rows, table["slug"], failures)
        if len(rows) != table["rows"]:
            fail(failures, f"{table['slug']}: manifest rows {table['rows']} != actual {len(rows)}")
        table_rows[table["slug"]] = rows
        total_rows += len(rows)
    if total_rows != EXPECTED["data_rows"]:
        fail(failures, f"data rows {total_rows} != {EXPECTED['data_rows']}")

    care_gaps = table_rows["care_gap_cells"]
    quality = table_rows["quality_measures"]
    if len(care_gaps) != EXPECTED["care_gap_cells"]:
        fail(failures, f"care gap cells {len(care_gaps)} != {EXPECTED['care_gap_cells']}")
    if len(quality) != EXPECTED["quality_measures"]:
        fail(failures, f"quality measures {len(quality)} != {EXPECTED['quality_measures']}")
    open_gap_total = sum(numeric(row, "open_gap_count") for row in care_gaps)
    closed_gap_total = sum(numeric(row, "closed_gap_count") for row in care_gaps)
    weighted_rate = round(closed_gap_total / (open_gap_total + closed_gap_total) * 100, 2)
    if open_gap_total != EXPECTED["open_care_gaps"]:
        fail(failures, f"open care gaps {open_gap_total} != {EXPECTED['open_care_gaps']}")
    if weighted_rate != EXPECTED["weighted_cell_closure_rate_pct"]:
        fail(failures, f"weighted closure {weighted_rate} != {EXPECTED['weighted_cell_closure_rate_pct']}")

    cohort_sums = defaultdict(int)
    measure_sums = defaultdict(int)
    for row in care_gaps:
        cohort_sums[row["cohort_key"]] += numeric(row, "open_gap_count")
        measure_sums[row["measure_code"]] += numeric(row, "open_gap_count")
    if len(measure_sums) != EXPECTED["quality_measures"]:
        fail(failures, f"care-gap measure coverage {len(measure_sums)} != {EXPECTED['quality_measures']}")
    measure_categories = {row["measure_code"]: row["domain"] for row in quality}
    for row in care_gaps:
        category = measure_categories[row["measure_code"]]
        cohort = row["cohort_key"]
        if row["measure_code"] in {"SUPD", "MAD", "MAH", "MAC"} and cohort not in {"ma_seniors", "medicare_snp"}:
            fail(failures, f"{row['measure_code']} incorrectly applies to {cohort}")
        if category == "child" and not cohort.startswith("pediatric_"):
            fail(failures, f"{row['measure_code']} child measure applies to non-pediatric cohort {cohort}")

    facilities = table_rows["facility_footprint"]
    attributed_lives = sum(numeric(row, "attributed_lives") for row in facilities)
    if attributed_lives != EXPECTED["attributed_lives"]:
        fail(failures, f"attributed lives {attributed_lives} != {EXPECTED['attributed_lives']}")

    interfaces = table_rows["interface_inventory"]
    source_counts = Counter(row["source_control_state"] for row in interfaces)
    monitor_counts = Counter(row["monitoring_state"] for row in interfaces)
    if source_counts["not_versioned"] != EXPECTED["unversioned_interfaces"]:
        fail(failures, f"unversioned interfaces {source_counts['not_versioned']} != {EXPECTED['unversioned_interfaces']}")
    if monitor_counts["not_monitored"] != EXPECTED["unmonitored_interfaces"]:
        fail(failures, f"unmonitored interfaces {monitor_counts['not_monitored']} != {EXPECTED['unmonitored_interfaces']}")

    org_roles = table_rows["org_roles"]
    open_reqs = [row for row in org_roles if row["vacancy_status"] == "open_requisition"]
    if len(open_reqs) != EXPECTED["org_open_requisitions"]:
        fail(failures, f"open requisitions {len(open_reqs)} != {EXPECTED['org_open_requisitions']}")
    vp_reports = [row for row in org_roles if row["reports_to_role"] == "VP Application Services"]
    if len(vp_reports) != 2:
        fail(failures, f"VP Application Services direct reports {len(vp_reports)} != 2")

    operating = table_rows["operating_evidence"]
    op_counts = Counter(row["record_type"] for row in operating)
    expected_operating = {
        "site": 48,
        "stakeholder": 34,
        "data_quality_finding": 30,
        "risk": 28,
        "time_motion_minute": 53,
        "evidence_artifact": 22,
        "constraint": 6,
        "phase_acceptance_check": 21,
        "prior_attempt": 3,
        "value_input": 3,
        "value_exclusion": 3,
        "upload_block": 5,
        "runbook_step": 26,
    }
    for key, expected in expected_operating.items():
        if op_counts[key] != expected:
            fail(failures, f"operating evidence {key} {op_counts[key]} != {expected}")
    search_minutes = sum(1 for row in operating if row["record_type"] == "time_motion_minute" and row["status"] == "search_or_reconciliation")
    if search_minutes != 31:
        fail(failures, f"time-and-motion search minutes {search_minutes} != 31")
    unvalidated_zero = [row for row in operating if row["record_type"] == "value_input" and row["status"] == "UNVALIDATED" and row["metric_value"] == "0"]
    if len(unvalidated_zero) != 3:
        fail(failures, f"UNVALIDATED zero value inputs {len(unvalidated_zero)} != 3")

    for doc in manifest["narrative_documents"]:
        validate_doc_pair(PACK_ROOT / doc["markdown"], PACK_ROOT / doc["docx"], failures)

    cheat = validate_cheat_sheet(failures)
    prohibited = scan_prohibited_strings(failures)
    pack_files = {path for path in PACK_ROOT.rglob("*") if path.is_file()}
    pack_files.update({REPORT_JSON, REPORT_MD})
    file_count = len(pack_files)
    if file_count != EXPECTED["file_count"]:
        fail(failures, f"file count {file_count} != {EXPECTED['file_count']}")

    report = {
        "ok": not failures,
        "failures": failures,
        "summary": {
            "file_count": file_count,
            "data_rows": total_rows,
            "tables": len(tables),
            "narratives": len(manifest["narrative_documents"]),
            "care_gap_cells": len(care_gaps),
            "quality_measures": len(quality),
            "attributed_lives": attributed_lives,
            "open_care_gaps": open_gap_total,
            "closed_care_gaps": closed_gap_total,
            "weighted_closure_rate_pct": weighted_rate,
            "unversioned_interfaces": source_counts["not_versioned"],
            "unmonitored_interfaces": monitor_counts["not_monitored"],
            "open_requisitions": len(open_reqs),
            "operating_evidence_counts": dict(sorted(op_counts.items())),
            "cheat_sheet": cheat,
            "prohibited_matches": prohibited,
        },
    }
    REPORT_JSON.write_text(json.dumps(report, indent=2) + "\n", encoding="utf-8")
    lines = [
        "# Population Health Command Center Fixture Validation",
        "",
        f"- Result: {'PASS' if report['ok'] else 'FAIL'}",
        f"- Files: {file_count}",
        f"- Data rows: {total_rows}",
        f"- Care-gap cells: {len(care_gaps)}",
        f"- Open care gaps: {open_gap_total:,}",
        f"- Weighted closure rate: {weighted_rate}%",
        f"- Cheat sheet: {cheat['phases']} phases, {cheat['fields']} fields, {cheat['upload_blocks']} upload blocks, {cheat['file_refs']} file refs",
        "",
    ]
    if failures:
        lines.append("## Failures")
        lines.extend(f"- {item}" for item in failures)
    else:
        lines.append("All validator checks passed.")
    REPORT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report


def main() -> None:
    report = validate_pack()
    print(json.dumps(report["summary"], indent=2))
    if not report["ok"]:
        print("\n".join(report["failures"]), file=sys.stderr)
        raise SystemExit(1)


if __name__ == "__main__":
    main()
